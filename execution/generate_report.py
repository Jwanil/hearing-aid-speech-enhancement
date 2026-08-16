"""
generate_report.py
------------------
Generates the Minor Project Report as a formatted .docx file.
Usage: python execution/generate_report.py

Output: docs/Minor_Project_Report_Hearing_Aid_Speech_Enhancement.docx

Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

# ─── Helpers ────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '4A90D9')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    for run in h.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xAA)
            run.font.size = Pt(16)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
            run.font.size = Pt(13)
        else:
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
            run.font.size = Pt(11)
    return h

def add_body(doc, text, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F5F5')
    pPr.append(shd)
    return p

def add_info_box(doc, text, color='EBF3FB'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.italic = True
    run.font.name = 'Calibri'
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)
    return p

# ─── Document ────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('MINOR PROJECT REPORT')
title_run.bold = True
title_run.font.size = Pt(13)
title_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
title_run.font.name = 'Calibri'

doc.add_paragraph()

main_title_p = doc.add_paragraph()
main_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
mt_run = main_title_p.add_run('Audiogram-Personalized, Low-Latency\nSpeech Enhancement for Hearing Aids')
mt_run.bold = True
mt_run.font.size = Pt(22)
mt_run.font.color.rgb = RGBColor(0x1A, 0x56, 0xAA)
mt_run.font.name = 'Calibri'

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub_p.add_run('Using Deep Learning and FiLM Conditioning\nfor Individualized Hearing Aid Enhancement')
sub_run.font.size = Pt(13)
sub_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
sub_run.font.italic = True
sub_run.font.name = 'Calibri'

add_horizontal_rule(doc)

doc.add_paragraph()

info_lines = [
    ('Submitted by:', 'Jwanil  |  Namya Shah'),
    ('Department:', '[Your Department Name]'),
    ('Institution:', '[Your College / University Name]'),
    ('Academic Year:', '2026 – 2027'),
    ('Submitted to:', '[Faculty Name], [Designation]'),
    ('Date:', datetime.date.today().strftime('%B %d, %Y')),
]

for label, value in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Abstract', 1)
add_horizontal_rule(doc)
add_body(doc,
    'Hearing aid users frequently cite the inability to understand speech in noisy environments as their primary '
    'complaint, despite advances in hearing aid technology. Current deep learning-based speech enhancement systems, '
    'while effective at noise reduction, treat all users identically, ignoring the fact that every individual '
    'has a unique hearing loss profile (audiogram). This project proposes a novel audiogram-personalized speech '
    'enhancement system that takes the user\'s individual audiogram as a runtime input and adapts its noise '
    'reduction behaviour accordingly, using Feature-wise Linear Modulation (FiLM) conditioning. '
    '\n\n'
    'We implement and compare three models: (1) a classical Wiener filter baseline, (2) a generic deep neural '
    'network (DNN) denoiser using a U-Net architecture with Ideal Ratio Mask prediction, and (3) our proposed '
    'audiogram-conditioned variant. Evaluation uses Hearing Aid Speech Perception Index (HASPI) and Hearing Aid '
    'Speech Quality Index (HASQI) — metrics specifically designed for hearing-impaired listeners — alongside '
    'standard STOI and SI-SDR measures. We additionally report model latency and parameter count, given the '
    'sub-10ms real-time processing constraint imposed by hearing aid hardware. '
    '\n\n'
    'We expect the personalized model to demonstrate measurable improvement in HASPI over the generic baseline, '
    'particularly for listeners with non-uniform hearing loss profiles (e.g., sloping high-frequency loss), '
    'establishing audiogram-conditioning as a practical and principled approach to individualized hearing aid '
    'enhancement.'
)

doc.add_paragraph()
add_info_box(doc,
    '🔑  Keywords: Speech Enhancement, Hearing Aids, FiLM Conditioning, Audiogram Personalization, '
    'Deep Learning, Ideal Ratio Mask, HASPI, HASQI, Low-Latency Inference'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, 'Table of Contents', 1)
add_horizontal_rule(doc)

toc_items = [
    ('1.', 'Introduction', ''),
    ('2.', 'Problem Statement', ''),
    ('3.', 'Literature Review', ''),
    ('   3.1', 'Classical Speech Enhancement', ''),
    ('   3.2', 'Deep Learning Approaches', ''),
    ('   3.3', 'Hearing-Aid-Specific Research', ''),
    ('   3.4', 'FiLM Conditioning', ''),
    ('4.', 'Background & Theory', ''),
    ('   4.1', 'Audio Signal Processing & Spectrograms', ''),
    ('   4.2', 'Hearing Loss & Audiograms', ''),
    ('   4.3', 'The Wiener Filter', ''),
    ('   4.4', 'Deep Neural Network Masking', ''),
    ('   4.5', 'FiLM: Feature-wise Linear Modulation', ''),
    ('5.', 'Proposed Methodology', ''),
    ('   5.1', 'System Architecture', ''),
    ('   5.2', 'Model 1: Wiener Filter Baseline', ''),
    ('   5.3', 'Model 2: Generic DNN Denoiser', ''),
    ('   5.4', 'Model 3: Audiogram-Personalized DNN', ''),
    ('   5.5', 'Training Strategy', ''),
    ('6.', 'Dataset & Tools', ''),
    ('7.', 'Evaluation Metrics', ''),
    ('8.', 'Project Timeline', ''),
    ('9.', 'Expected Results & Contributions', ''),
    ('10.', 'Limitations & Future Work', ''),
    ('11.', 'Conclusion', ''),
    ('12.', 'References', ''),
]

for num, title, _ in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.1 if num.startswith('   ') else 0)
    r = p.add_run(f'{num}    {title}')
    r.font.size = Pt(10.5)
    r.font.name = 'Calibri'
    r.bold = not num.startswith('   ')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '1. Introduction', 1)
add_horizontal_rule(doc)
add_body(doc,
    'Hearing loss is one of the most prevalent sensory impairments worldwide. According to the World Health '
    'Organization (WHO), approximately 466 million people — or roughly 5% of the global population — have '
    'disabling hearing loss, a figure projected to rise to 900 million by 2050. For these individuals, hearing '
    'aids are the primary assistive technology, amplifying acoustic signals to compensate for reduced auditory '
    'sensitivity.'
)
add_body(doc,
    'Despite significant advances in digital signal processing and, more recently, deep learning, the most '
    'persistent complaint among hearing aid users remains difficulty understanding speech in noisy environments. '
    'The "cocktail party problem" — segregating a target speaker from competing background noise and other '
    'speakers — remains deeply challenging for hearing-impaired listeners, whose damaged cochlear hair cells '
    'further degrade the frequency resolution required for such segregation.'
)
add_body(doc,
    'Modern hearing aids address this through programmable amplification that is personalised to the user\'s '
    'audiogram — a clinically measured chart of their hearing threshold at each frequency. However, the noise '
    'reduction algorithms applied upstream of this amplification stage remain generic: identical processing is '
    'applied to every user, regardless of their individual hearing loss profile. This represents a fundamental '
    'mismatch, since the frequencies most critical to intelligibility differ across individuals.'
)
add_body(doc,
    'This project proposes closing this gap by conditioning a deep learning speech enhancement model on the '
    'user\'s audiogram at inference time, enabling truly personalised noise reduction that adapts to the specific '
    'pattern of each user\'s hearing loss.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '2. Problem Statement', 1)
add_horizontal_rule(doc)
add_body(doc,
    'Given a noisy speech signal y(t) = s(t) + n(t), where s(t) is clean speech and n(t) is additive noise, '
    'and given a user-specific audiogram vector a ∈ ℝ⁶ representing hearing thresholds at standard audiometric '
    'frequencies {250, 500, 1000, 2000, 4000, 8000} Hz, the goal is to learn a mapping:'
)
add_code_block(doc, '    f(Y, a) → Ŝ\n\n    where Y is the STFT of y(t) and Ŝ approximates the STFT of s(t)')
add_body(doc,
    'such that the quality and intelligibility of the reconstructed speech, as assessed by a hearing-loss-aware '
    'metric (HASPI, HASQI), is maximised for a listener with audiogram a. Critically, f must adapt its '
    'processing based on a — two users with different audiograms must receive different enhancements for the '
    'same input signal.'
)

add_heading(doc, 'Key Constraints', 2)
add_bullet(doc, 'Latency: hearing aids require < 10ms processing delay to avoid perceptible lip-sync mismatch.')
add_bullet(doc, 'On-device compute: real HA chips operate with tight memory and FLOP budgets (~1M parameters).')
add_bullet(doc, 'No retraining per user: the audiogram is a runtime input, not a training-time choice.')
add_bullet(doc, 'No clinical subjects required: evaluation uses validated objective metrics (HASPI/HASQI).')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. LITERATURE REVIEW
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '3. Literature Review', 1)
add_horizontal_rule(doc)

add_heading(doc, '3.1 Classical Speech Enhancement', 2)
add_body(doc,
    'Classical approaches include spectral subtraction (Boll, 1979), which estimates and subtracts a noise '
    'spectrum but produces musical noise artefacts, and the Wiener filter (Wiener, 1949), which computes an '
    'optimal linear filter minimising mean-squared error between the enhanced and clean signal. These methods '
    'assume stationary noise — an assumption that fails in real-world environments. They serve as our '
    'comparison baseline.'
)

add_heading(doc, '3.2 Deep Learning Approaches', 2)
add_body(doc,
    'The introduction of Deep Neural Networks (DNNs) for speech enhancement marked a paradigm shift. '
    'Xu et al. (2014) demonstrated that DNNs could learn speech-to-noise mapping directly from data. '
    'Subsequent work introduced mask-based approaches — predicting an Ideal Binary Mask (IBM) or Ideal '
    'Ratio Mask (IRM) in the time-frequency domain — which preserve the original signal structure while '
    'suppressing noise.'
)
add_body(doc,
    'Encoder-decoder architectures such as U-Net (Ronneberger et al., 2015), originally designed for image '
    'segmentation, found natural application in spectrogram masking. Conv-TasNet (Luo & Mesgarani, 2019) '
    'demonstrated that operating directly in the time domain with depthwise-separable convolutions achieves '
    'state-of-the-art performance with millisecond-scale latency, directly relevant to the hearing aid constraint.'
)

add_heading(doc, '3.3 Hearing-Aid-Specific Research', 2)
add_body(doc,
    'Diehl et al. (2023), published in Nature Scientific Reports, demonstrated that a U-Net model trained '
    'on a large proprietary dataset could restore speech intelligibility for hearing-impaired users to '
    'near-normal-hearing levels on standard benchmarks. This represents the strongest existing result in '
    'generic DNN denoising for hearing aids and is the primary system our work differentiates from.'
)
add_body(doc,
    'The Clarity Enhancement Challenge (University of Salford / CVSSP, 2021-2024) established standardised '
    'benchmarks, datasets with real listener audiograms, and hearing-loss-aware evaluation metrics (HASPI, '
    'HASQI) for this exact research problem. NeuroAMP (2025) proposes a personalised neural amplifier '
    'conditioned on audiograms — but for the amplification stage, not the noise reduction stage.'
)
add_info_box(doc,
    '📌  Gap: No published system combines audiogram-conditioned noise reduction with explicit latency-awareness '
    'and evaluation using HASPI/HASQI across diverse individual hearing profiles. This is the specific gap '
    'our project targets.'
)

add_heading(doc, '3.4 FiLM Conditioning', 2)
add_body(doc,
    'Feature-wise Linear Modulation (FiLM) was introduced by Perez et al. (2017) for visual question '
    'answering — injecting linguistic conditioning into a vision network. It has since been applied to '
    'speech synthesis (Arik et al., 2018) and multi-speaker adaptation. The mechanism is elegant and '
    'computationally cheap: a small conditioning network generates per-channel scale (γ) and shift (β) '
    'parameters that modulate intermediate feature maps. We adopt this as our personalisation mechanism.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. BACKGROUND & THEORY
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '4. Background & Theory', 1)
add_horizontal_rule(doc)

add_heading(doc, '4.1 Audio Signal Processing & Spectrograms', 2)
add_body(doc,
    'A digital audio signal is a sequence of amplitude samples at a fixed sampling rate (16,000 Hz for speech '
    'processing). The Short-Time Fourier Transform (STFT) converts this one-dimensional signal into a '
    'two-dimensional time-frequency representation (spectrogram) by applying the Discrete Fourier Transform '
    'to short overlapping windows (typically 25ms, stride 10ms):'
)
add_code_block(doc, '    X[k, n] = Σ x[m] · w[n - m] · e^(-j2πkm/N)\n\n    where w is the analysis window, k is frequency bin, n is frame index')
add_body(doc,
    'The STFT produces complex-valued coefficients. We work with the magnitude spectrogram |X[k,n]| and '
    'retain the phase from the noisy signal for reconstruction via the Inverse STFT (ISTFT). This '
    'simplification — the noisy phase approximation — is standard in mask-based enhancement and has '
    'negligible perceptual impact at moderate SNR.'
)

add_heading(doc, '4.2 Hearing Loss & Audiograms', 2)
add_body(doc,
    'Sensorineural hearing loss (SNHL) — the most common form — results from damage to cochlear hair cells. '
    'These cells transduce mechanical vibrations into neural signals and are tonotopically organised: '
    'high-frequency cells at the basal end, low-frequency at the apical end. Noise exposure and ageing '
    'predominantly damage basal (high-frequency) cells first.'
)
add_body(doc,
    'An audiogram maps hearing threshold (dB Hearing Level, dB HL) against frequency (250–8000 Hz). '
    'A threshold of 0 dB HL indicates normal sensitivity; a threshold of 60 dB HL means a tone must be '
    '60 dB louder than normal to be detected. Clinically, loss is classified as:'
)
# Build table
table = doc.add_table(rows=6, cols=2)
table.style = 'Table Grid'
headers = ['Classification', 'Threshold Range (dB HL)']
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    set_cell_bg(cell, '1A56AA')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

rows_data = [
    ('Normal', '0 – 25'),
    ('Mild Loss', '26 – 40'),
    ('Moderate Loss', '41 – 55'),
    ('Moderately Severe', '56 – 70'),
    ('Severe Loss', '71 – 90'),
]
for i, (label, val) in enumerate(rows_data, 1):
    table.cell(i, 0).text = label
    table.cell(i, 1).text = val
    if i % 2 == 0:
        set_cell_bg(table.cell(i, 0), 'EBF3FB')
        set_cell_bg(table.cell(i, 1), 'EBF3FB')
    for cell in [table.cell(i, 0), table.cell(i, 1)]:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_body(doc,
    'For our model, the audiogram is represented as a 6-dimensional vector of threshold values at standard '
    'frequencies: a = [a₂₅₀, a₅₀₀, a₁₀₀₀, a₂₀₀₀, a₄₀₀₀, a₈₀₀₀] ∈ ℝ⁶, normalised to the range [0, 1].'
)

add_heading(doc, '4.3 The Wiener Filter', 2)
add_body(doc,
    'The Wiener filter computes a frequency-dependent gain W(k,n) that minimises the mean squared error '
    'between the filtered output and the clean signal. Under the assumption of uncorrelated speech and noise:'
)
add_code_block(doc, '    W(k,n) = λ_s(k,n) / [λ_s(k,n) + λ_n(k,n)]\n\n    λ_s = speech power estimate, λ_n = noise power estimate')
add_body(doc,
    'When speech power dominates (high SNR), W ≈ 1 (pass through). When noise dominates (low SNR), W ≈ 0 '
    '(suppress). The enhanced spectrogram is Ŷ = W · Y. The challenge lies in estimating λ_s and λ_n '
    'from the noisy observation alone. Classical implementations assume noise stationarity — an assumption '
    'violated in real environments.'
)

add_heading(doc, '4.4 Deep Neural Network Masking', 2)
add_body(doc,
    'Rather than estimating noise statistics analytically, a DNN learns to predict the mask M(k,n) directly '
    'from data. The Ideal Ratio Mask (IRM) is defined as:'
)
add_code_block(doc, '    IRM(k,n) = |S(k,n)|² / [|S(k,n)|² + |N(k,n)|²]')
add_body(doc,
    'A U-Net with sigmoid output predicts M ≈ IRM. The enhanced spectrogram is Ŷ = M · Y. This approach '
    'generalises to non-stationary noise and complex acoustic environments because it learns from data '
    'rather than hand-crafted statistical models.'
)

add_heading(doc, '4.5 FiLM: Feature-wise Linear Modulation', 2)
add_body(doc,
    'FiLM injects conditioning information (the audiogram vector a) into an intermediate layer of the main '
    'network. A small auxiliary network F (the FiLM generator) maps a to per-channel scale and shift parameters:'
)
add_code_block(doc, '    [γ, β] = F(a)    where γ, β ∈ ℝᶜ  (C = number of feature channels)\n\n    FiLM(h) = γ ⊙ h + β\n\n    where h is the feature map tensor and ⊙ is element-wise multiplication')
add_body(doc,
    'The FiLM generator F is jointly trained with the main U-Net. The model learns, end-to-end, which '
    'internal features to amplify or suppress for each audiogram profile — no hand-crafted rules required. '
    'The computational overhead is minimal: F adds only ~C × (dim_a + 1) parameters.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROPOSED METHODOLOGY
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '5. Proposed Methodology', 1)
add_horizontal_rule(doc)

add_heading(doc, '5.1 System Architecture', 2)
add_code_block(doc,
    '  Noisy speech y(t)\n'
    '       │\n'
    '       ▼\n'
    '  ┌──────────┐\n'
    '  │   STFT   │  → magnitude spectrogram Y ∈ ℝ^{F×T}\n'
    '  └────┬─────┘\n'
    '       │\n'
    '       ▼\n'
    '  ┌─────────────────┐\n'
    '  │  Encoder (CNN)  │  Conv layers: F×T → C×(F/k)×(T/k)\n'
    '  │  Layers 1–4     │  (downsampling via strided conv)\n'
    '  └────────┬────────┘\n'
    '           │\n'
    '           ▼\n'
    '  ┌────────────────────────────┐    ┌─────────────────────┐\n'
    '  │    FiLM Conditioning       │◄───│  Audiogram vector   │\n'
    '  │    h = γ(a) ⊙ h + β(a)   │    │  a ∈ ℝ⁶            │\n'
    '  └────────┬───────────────────┘    └─────────────────────┘\n'
    '           │\n'
    '           ▼\n'
    '  ┌─────────────────┐\n'
    '  │  Decoder (CNN)  │  Layers 4–1 + skip connections\n'
    '  └────────┬────────┘\n'
    '           │\n'
    '           ▼\n'
    '  ┌──────────────────┐\n'
    '  │  Mask M ∈ [0,1]  │  Sigmoid activation\n'
    '  └────────┬─────────┘\n'
    '           │\n'
    '      M ⊙ Y  →  Ŷ (enhanced spectrogram)\n'
    '           │\n'
    '           ▼\n'
    '  ┌──────────┐\n'
    '  │  ISTFT   │  → enhanced speech ŝ(t)\n'
    '  └──────────┘'
)

add_heading(doc, '5.2 Model 1: Wiener Filter Baseline', 2)
add_bullet(doc, 'Type: Classical DSP (no learning)')
add_bullet(doc, 'Implementation: NumPy/SciPy spectral estimation with a priori SNR update (decision-directed approach)')
add_bullet(doc, 'Noise estimate: minimum statistics tracking over 1.5s window')
add_bullet(doc, 'Purpose: Establishes a well-understood quantitative baseline')

add_heading(doc, '5.3 Model 2: Generic DNN Denoiser', 2)
add_bullet(doc, 'Architecture: U-Net with 4 encoder and 4 decoder stages, skip connections, BatchNorm + ReLU')
add_bullet(doc, 'Input: Noisy magnitude spectrogram Y ∈ ℝ^{257×T} (N_FFT=512, 16kHz)')
add_bullet(doc, 'Output: Mask M ∈ [0,1]^{257×T} (Sigmoid activation)')
add_bullet(doc, 'Loss: SI-SDR on reconstructed waveform')
add_bullet(doc, 'Audiogram: Not used. Same processing for every user.')

add_heading(doc, '5.4 Model 3: Audiogram-Personalized DNN (Proposed)', 2)
add_bullet(doc, 'Architecture: Identical U-Net + FiLM layer at bottleneck')
add_bullet(doc, 'Additional input: Normalised audiogram vector a ∈ [0,1]⁶')
add_bullet(doc, 'FiLM generator: 2-layer MLP (6 → 128 → 2C), outputs [γ, β] ∈ ℝ^{2C}')
add_bullet(doc, 'Applied to: encoder bottleneck feature maps (64 channels)')
add_bullet(doc, 'Parameter overhead: ~16,640 additional parameters vs generic model')
add_bullet(doc, 'Loss: Same SI-SDR loss; audiogram conditions behaviour, not objective')

add_heading(doc, '5.5 Training Strategy', 2)
add_bullet(doc, 'Optimiser: Adam, lr=3×10⁻⁴, cosine annealing schedule')
add_bullet(doc, 'Batch size: 16')
add_bullet(doc, 'Training signal: on-the-fly mixing at random SNR ∈ [-5, 10] dB')
add_bullet(doc, 'Audiogram generation: random synthetic profiles via NAL-R distribution (mild/moderate/severe, flat/sloping/reverse-sloping)')
add_bullet(doc, 'Epochs: 100 (with early stopping on HASPI validation score)')
add_bullet(doc, 'Hardware: GPU if available; CPU fallback for development')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. DATASETS & TOOLS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '6. Dataset & Tools', 1)
add_horizontal_rule(doc)

add_heading(doc, 'Datasets', 2)
table2 = doc.add_table(rows=3, cols=3)
table2.style = 'Table Grid'
for i, h in enumerate(['Dataset', 'Purpose', 'Access']):
    cell = table2.cell(0, i)
    cell.text = h
    set_cell_bg(cell, '1A56AA')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

data2 = [
    ('Clarity Challenge\n(CEC2 / CEC3)', 'Primary — purpose-built for HA research. Includes listener audiograms, room acoustics, HA-specific evaluation. The definitive benchmark.', 'claritychallenge.org\n(free, registration)'),
    ('VoiceBank-DEMAND', 'Standard noisy speech benchmark for model sanity checking and pretraining the generic baseline.', 'Available via SpeechBrain'),
]
for i, row in enumerate(data2, 1):
    for j, val in enumerate(row):
        table2.cell(i, j).text = val
        if i % 2 == 0:
            set_cell_bg(table2.cell(i, j), 'EBF3FB')
        for p in table2.cell(i, j).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
doc.add_paragraph()

add_heading(doc, 'Software Stack', 2)
table3 = doc.add_table(rows=8, cols=2)
table3.style = 'Table Grid'
for i, h in enumerate(['Library / Tool', 'Role in Project']):
    cell = table3.cell(0, i)
    cell.text = h
    set_cell_bg(cell, '1A56AA')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

data3 = [
    ('PyTorch + torchaudio', 'Model definition, training loop, STFT/ISTFT, audio I/O'),
    ('SpeechBrain', 'Pretrained enhancement baselines, training recipe reference'),
    ('Asteroid', 'Conv-TasNet reference implementation and SI-SDR loss'),
    ('pyclarity', 'HASPI/HASQI computation and hearing loss simulation (cochlear model)'),
    ('pystoi', 'STOI/ESTOI computation'),
    ('librosa + matplotlib', 'Spectrogram visualisation and audio analysis'),
    ('TensorBoard', 'Training metrics visualisation'),
]
for i, (tool, role) in enumerate(data3, 1):
    table3.cell(i, 0).text = tool
    table3.cell(i, 1).text = role
    if i % 2 == 0:
        set_cell_bg(table3.cell(i, 0), 'EBF3FB')
        set_cell_bg(table3.cell(i, 1), 'EBF3FB')
    for col in [0, 1]:
        for p in table3.cell(i, col).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. EVALUATION METRICS
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '7. Evaluation Metrics', 1)
add_horizontal_rule(doc)
add_body(doc,
    'Standard audio quality metrics (PESQ, SNR) do not account for the specific perceptual characteristics '
    'of hearing-impaired listeners. We use the following metric suite:'
)

table4 = doc.add_table(rows=6, cols=3)
table4.style = 'Table Grid'
for i, h in enumerate(['Metric', 'Range', 'What It Measures']):
    cell = table4.cell(0, i)
    cell.text = h
    set_cell_bg(cell, '1A56AA')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

data4 = [
    ('HASPI ↑', '0 – 1', 'Hearing Aid Speech Perception Index. Intelligibility for hearing-impaired listeners. Takes audiogram as input. Primary metric.'),
    ('HASQI ↑', '0 – 1', 'Hearing Aid Speech Quality Index. Perceived quality for hearing-impaired listeners. Takes audiogram as input.'),
    ('STOI / ESTOI ↑', '0 – 1', 'Short-Time Objective Intelligibility. Standard baseline; does NOT account for hearing loss.'),
    ('SI-SDR ↑', 'dB', 'Scale-Invariant Signal-to-Distortion Ratio. Overall signal quality.'),
    ('Latency ↓', 'ms', 'Inference time per frame. Real-time threshold for hearing aids: < 10ms.'),
]
for i, row in enumerate(data4, 1):
    for j, val in enumerate(row):
        table4.cell(i, j).text = val
        if i % 2 == 0:
            set_cell_bg(table4.cell(i, j), 'EBF3FB')
        for p in table4.cell(i, j).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)

doc.add_paragraph()
add_info_box(doc,
    '📌  Key insight: Improvement in HASPI (hearing-loss-aware) should be larger than improvement in STOI '
    '(hearing-loss-unaware) for the personalised model. This is the expected signature of successful '
    'audiogram conditioning — the model has learned to prioritise the frequencies most critical for that '
    'specific listener\'s hearing profile.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. PROJECT TIMELINE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, '8. Project Timeline', 1)
add_horizontal_rule(doc)

table5 = doc.add_table(rows=9, cols=4)
table5.style = 'Table Grid'
for i, h in enumerate(['Phase', 'Task', 'Weeks', 'Responsible']):
    cell = table5.cell(0, i)
    cell.text = h
    set_cell_bg(cell, '1A56AA')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)

timeline_data = [
    ('0', 'Audio/DSP Fundamentals — STFT, spectrograms, hearing loss simulation', '1–2', 'Both'),
    ('1', 'Audiology basics — Audiogram generation, pyclarity hearing loss sim', '2', 'Both'),
    ('2', 'Classical Wiener Filter baseline implementation and evaluation', '3', 'Jwanil'),
    ('3', 'PyTorch data pipeline — noisy mixing, DataLoader, batching', '4', 'Namya'),
    ('4', 'Generic DNN denoiser — U-Net training on VoiceBank-DEMAND', '5–7', 'Both'),
    ('5', 'FiLM conditioning — Personalised DNN, retrain on Clarity dataset', '8–10', 'Both'),
    ('6', 'Full evaluation — HASPI/HASQI, latency, parameter count, plots', '11', 'Both'),
    ('7', 'Report writing, audio demos, final presentation preparation', '12+', 'Both'),
]
for i, row in enumerate(timeline_data, 1):
    for j, val in enumerate(row):
        table5.cell(i, j).text = val
        if i % 2 == 0:
            set_cell_bg(table5.cell(i, j), 'EBF3FB')
        for p in table5.cell(i, j).paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)

# ══════════════════════════════════════════════════════════════════════════════
# 9. EXPECTED RESULTS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, '9. Expected Results & Contributions', 1)
add_horizontal_rule(doc)

add_heading(doc, 'Quantitative Expectations', 2)
add_body(doc,
    'Based on the literature and the nature of the contribution, we expect:'
)
add_bullet(doc, 'Personalised DNN > Generic DNN > Wiener Filter on HASPI and HASQI across all test audiograms.')
add_bullet(doc, 'The HASPI gap between personalised and generic should be largest for extreme audiograms (e.g., severe high-frequency loss), where the conditioning signal carries the most information.')
add_bullet(doc, 'STOI improvements should be similar between personalised and generic (since STOI ignores hearing loss), validating that the improvement is specifically due to personalisation, not just better noise reduction.')
add_bullet(doc, 'Latency: model inference should be reportable in absolute milliseconds for honest comparison against the 10ms constraint.')

add_heading(doc, 'Deliverables', 2)
add_numbered(doc, 'Three trained model variants (Wiener / Generic DNN / Personalised DNN) with saved weights.')
add_numbered(doc, 'Evaluation report: HASPI, HASQI, STOI, SI-SDR, latency, parameter count in tabular form.')
add_numbered(doc, 'Before/after audio samples for 3 audiogram profiles (mild, moderate, severe sloping).')
add_numbered(doc, 'Python scripts (execution/) for full reproducibility — any reader can retrain from scratch.')
add_numbered(doc, 'This written report.')

add_heading(doc, 'Novel Contributions', 2)
add_bullet(doc, 'FiLM conditioning applied to hearing aid speech enhancement (not previously published in combination with HASPI/HASQI evaluation).')
add_bullet(doc, 'Explicit audiogram-to-noise-mask conditioning without per-user retraining.')
add_bullet(doc, 'Latency and model-size analysis alongside intelligibility metrics — a rarely reported combination in academic HA research.')

# ══════════════════════════════════════════════════════════════════════════════
# 10. LIMITATIONS & FUTURE WORK
# ══════════════════════════════════════════════════════════════════════════════

add_heading(doc, '10. Limitations & Future Work', 1)
add_horizontal_rule(doc)

add_heading(doc, 'Limitations', 2)
add_bullet(doc, 'Hardware deployment: we do not deploy on real hearing aid chips. Latency is measured on CPU as a proxy.')
add_bullet(doc, 'Synthetic audiograms: training audiograms are synthetically generated. Real clinical audiograms may have different distributional properties.')
add_bullet(doc, 'No subjective listening tests: HASPI/HASQI are validated surrogates but cannot fully replace clinical trials with hearing-impaired subjects.')
add_bullet(doc, 'English-only dataset: Clarity Challenge and VoiceBank-DEMAND use English speech. Indian language phoneme sets may differ.')

add_heading(doc, 'Future Work', 2)
add_bullet(doc, 'Model compression: quantisation and pruning to hit hearing-aid-grade compute budgets.')
add_bullet(doc, 'Self-calibration: replace clinically measured audiograms with a short adaptive preference test (active learning).')
add_bullet(doc, 'Indian language evaluation: test on Hindi/Gujarati speech-in-noise (genuinely under-explored in HA research).')
add_bullet(doc, 'Multi-microphone extension: leverage beamforming for spatial noise reduction upstream of the DNN.')

# ══════════════════════════════════════════════════════════════════════════════
# 11. CONCLUSION
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, '11. Conclusion', 1)
add_horizontal_rule(doc)
add_body(doc,
    'This project proposes a targeted solution to a well-defined gap in hearing aid technology: while '
    'personalised amplification is standard practice, personalised noise reduction is not. By conditioning '
    'a deep learning speech enhancement model on the user\'s individual audiogram using FiLM, we enable '
    'the model to adapt its masking behaviour per person — preserving the speech cues that matter most '
    'for that specific listener\'s hearing profile.'
)
add_body(doc,
    'The three-model comparison (Wiener filter → Generic DNN → Personalised DNN), evaluated with '
    'hearing-loss-aware metrics (HASPI, HASQI) and supplemented by latency analysis, provides a '
    'rigorous and honest assessment of the contribution. The project is grounded in established '
    'benchmarks (Clarity Challenge), uses open-source tools (pyclarity, SpeechBrain), and is designed '
    'to be fully reproducible.'
)
add_body(doc,
    'If the personalised model achieves even a modest HASPI improvement over the generic baseline — '
    'particularly for severe or non-uniform audiogram profiles — it validates the core hypothesis and '
    'opens a practical path towards truly individualised hearing aid processing, a meaningful step towards '
    'improved quality of life for 466 million people worldwide.'
)

# ══════════════════════════════════════════════════════════════════════════════
# 12. REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading(doc, '12. References', 1)
add_horizontal_rule(doc)

references = [
    '[1]  Perez, E., Strub, F., de Vries, H., Dumoulin, V., & Courville, A. (2017). FiLM: Visual Reasoning with a General Conditioning Layer. AAAI 2018.',
    '[2]  Diehl, P., et al. (2023). Restoring speech intelligibility for hearing aid users with deep learning. Nature Scientific Reports, 13, 870.',
    '[3]  Luo, Y., & Mesgarani, N. (2019). Conv-TasNet: Surpassing Ideal Time-Frequency Magnitude Masking for Speech Separation. IEEE/ACM Transactions on Audio, Speech, and Language Processing.',
    '[4]  Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI 2015.',
    '[5]  Kates, J.M., & Arehart, K.H. (2014). The Hearing-Aid Speech Quality Index (HASQI) Version 2. Journal of the Audio Engineering Society, 62(3).',
    '[6]  Kates, J.M., & Arehart, K.H. (2021). The Hearing-Aid Speech Perception Index (HASPI) Version 2. Speech Communication, 131, 35-46.',
    '[7]  Clarity Enhancement Challenge Team (2021-2024). The Clarity Enhancement Challenge. University of Salford / CVSSP. https://claritychallenge.org',
    '[8]  NeuroAMP: A Novel End-to-end General Purpose Deep Neural Amplifier for Personalized Hearing Aids. arXiv:2502.10822 (2025).',
    '[9]  Xu, Y., Du, J., Dai, L.-R., & Lee, C.-H. (2014). A Regression Approach to Speech Enhancement Based on Deep Neural Networks. IEEE/ACM TASLP.',
    '[10] Loizou, P.C. (2007). Speech Enhancement: Theory and Practice. CRC Press.',
    '[11] Jensen, J., & Taal, C.H. (2016). An Algorithm for Predicting the Intelligibility of Speech Masked by Modulated Noise Maskers. IEEE/ACM TASLP.',
    '[12] World Health Organization (2021). World report on hearing. WHO Press, Geneva.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    run = p.add_run(ref)
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'

# ─── Save ────────────────────────────────────────────────────────────────────

output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'docs',
    'Minor_Project_Report_Hearing_Aid_Speech_Enhancement.docx'
)

doc.save(output_path)
print(f'\n✅  Report saved to:\n    {output_path}\n')
