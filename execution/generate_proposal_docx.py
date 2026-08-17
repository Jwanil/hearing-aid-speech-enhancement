"""
generate_proposal_docx.py
Generates the Minor Project PROPOSAL as a .docx — SAME design as generate_report.py.
White background | 1A56AA blue headings | EBF3FB alternating rows | F5F5F5 code blocks

Usage : python execution/generate_proposal_docx.py
Output: docs/Minor_Project_Proposal_Hearing_Aid_Speech_Enhancement.docx
Requires: pip install python-docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

BLUE_DARK = RGBColor(0x1A, 0x56, 0xAA)
BLUE_MID  = RGBColor(0x2E, 0x75, 0xB6)
GREY_DARK = RGBColor(0x40, 0x40, 0x40)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
HDR_HEX   = '1A56AA'
ALT_HEX   = 'EBF3FB'
RULE_HEX  = '4A90D9'
BOX_HEX   = 'EBF3FB'
BOX_TIP   = 'E8F8F0'
CODE_HEX  = 'F5F5F5'

def add_hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    btm = OxmlElement('w:bottom')
    btm.set(qn('w:val'), 'single'); btm.set(qn('w:sz'), '6')
    btm.set(qn('w:space'), '1'); btm.set(qn('w:color'), RULE_HEX)
    pBdr.append(btm); pPr.append(pBdr)

def set_bg(cell, hx):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hx)
    tcPr.append(shd)

def hd(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level==1 else 10)
    h.paragraph_format.space_after  = Pt(6)
    for r in h.runs:
        r.font.color.rgb = [BLUE_DARK, BLUE_MID, GREY_DARK][min(level-1,2)]
        r.font.size = [Pt(16), Pt(13), Pt(11)][min(level-1,2)]
    return h

def body(doc, text, bold=False, italic=False, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa); p.paragraph_format.space_before = Pt(2)
    r = p.add_run(text); r.font.size = Pt(11); r.font.bold = bold
    r.font.italic = italic; r.font.name = 'Calibri'; return p

def blt(doc, text, lvl=0, bp=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25 + lvl*0.25)
    if bp:
        r0 = p.add_run(bp+': '); r0.bold=True; r0.font.size=Pt(10.5); r0.font.name='Calibri'
    r = p.add_run(text); r.font.size = Pt(10.5); r.font.name = 'Calibri'; return p

def num(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.size=Pt(10.5); r.font.name='Calibri'; return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.left_indent=Inches(0.3)
    r = p.add_run(text); r.font.name='Courier New'; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(0x1E,0x1E,0x1E)
    pPr = p._p.get_or_add_pPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),CODE_HEX)
    pPr.append(shd); return p

def box(doc, text, color=BOX_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(6)
    p.paragraph_format.left_indent=Inches(0.2); p.paragraph_format.right_indent=Inches(0.2)
    r = p.add_run(text); r.font.size=Pt(10.5); r.font.italic=True; r.font.name='Calibri'
    pPr = p._p.get_or_add_pPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),color)
    pPr.append(shd); return p

def tbl(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i,h in enumerate(headers):
        c = t.cell(0,i); set_bg(c, HDR_HEX); c.text = h
        r = c.paragraphs[0].runs[0]; r.font.color.rgb=WHITE; r.bold=True; r.font.size=Pt(10)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.cell(ri+1,ci)
            if ri%2==0: set_bg(c, ALT_HEX)
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9.5); r.font.name='Calibri'
    doc.add_paragraph(); return t

# ── Document Setup ─────────────────────────────────────────────────────────
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2.5); s.bottom_margin=Cm(2.5); s.left_margin=Cm(3.0); s.right_margin=Cm(2.5)
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(11)

# ── TITLE PAGE ─────────────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('MINOR PROJECT PROPOSAL'); r.bold=True; r.font.size=Pt(13)
r.font.color.rgb=RGBColor(0x55,0x55,0x55); r.font.name='Calibri'
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Audiogram-Personalized, Low-Latency\nSpeech Enhancement for Hearing Aids')
r.bold=True; r.font.size=Pt(22); r.font.color.rgb=BLUE_DARK; r.font.name='Calibri'
doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Using Deep Learning, FiLM Conditioning, and Mamba State Space Models\nfor Individualized, Low-Latency Hearing Aid Enhancement')
r.font.size=Pt(13); r.font.italic=True; r.font.color.rgb=GREY_DARK; r.font.name='Calibri'
add_hr(doc); doc.add_paragraph()
for label, value in [
    ('Submitted by:','Jwanil Modi (23BIT194)  |  Namya Shah (23BIT027)'),
    ('Department:','Department of Information Technology'),
    ('Institution:','[Your College / University Name]'),
    ('Academic Year:','2026 – 2027'),
    ('Submitted to:','[Faculty Name], [Designation]'),
    ('Date:', datetime.date.today().strftime('%B %d, %Y')),
    ('Deadline:','November 1, 2026'),
]:
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r1=p.add_run(f'{label}  '); r1.bold=True; r1.font.size=Pt(11); r1.font.name='Calibri'
    r2=p.add_run(value); r2.font.size=Pt(11); r2.font.name='Calibri'
    p.paragraph_format.space_after=Pt(4)
doc.add_page_break()

# ── ABSTRACT ───────────────────────────────────────────────────────────────
hd(doc,'Abstract'); add_hr(doc)
body(doc,
'Hearing aid users consistently cite difficulty understanding speech in noisy environments as their '
'primary complaint. Existing deep learning-based speech enhancement systems achieve impressive noise '
'reduction, but treat all users identically — ignoring that each individual has a unique hearing loss '
'profile (audiogram). This project proposes a 5-model comparison study culminating in a novel '
'audiogram-personalized speech enhancement system that uses Feature-wise Linear Modulation (FiLM) '
'conditioning to adapt noise-reduction behaviour per user.\n\n'
'We implement and compare: (1) Wavelet DWT (Discrete Wavelet Transform) denoising, (2) MMSE-LSA '
'(Minimum Mean Square Error — Log Spectral Amplitude) filtering — the current industry standard in '
'commercial hearing aids, (3) 1D CNN (Conv-TasNet style) waveform model, (4) U-Net with '
'Self-Attention + FiLM + CRM (Complex Ratio Mask) + MetricGAN+ trained with a HASPI discriminator, '
'and (5) Mamba/SSM (State Space Model) + FiLM — current SOTA (PESQ 3.69). Evaluation uses HASPI '
'(Hearing Aid Speech Perception Index) and HASQI (Hearing Aid Speech Quality Index) — validated for '
'hearing-impaired listeners — alongside STOI, PESQ, SI-SDR, and latency.\n\n'
'Core novelty: training a MetricGAN+ discriminator on HASPI rather than PESQ, and applying FiLM to '
'a Mamba architecture for hearing aids — neither is published. Expected outcome: ΔHASPI >> ΔSTOI for '
'personalized models vs generic model, proving audiogram conditioning provides hearing-specific benefit.')
doc.add_paragraph()
box(doc,'Keywords: Speech Enhancement, Hearing Aids, Audiogram Personalization, FiLM Conditioning, '
'Mamba SSM, Complex Ratio Mask, HASPI, MetricGAN+, Wavelet Denoising, MMSE-LSA, Low-Latency Inference')
doc.add_page_break()

# ── TABLE OF CONTENTS ──────────────────────────────────────────────────────
hd(doc,'Table of Contents'); add_hr(doc)
toc_items = [
    ('1.','Introduction'), ('2.','Problem Statement'),
    ('3.','Literature Review & Research Gap'),
    ('   3.1','Classical Speech Enhancement'), ('   3.2','Deep Learning Approaches'),
    ('   3.3','Hearing-Aid-Specific Research'), ('   3.4','Faculty Feedback Integration'),
    ('   3.5','Post-Research-Sweep Architecture Upgrades'),
    ('4.','Background & Theory'),
    ('   4.1','Audio Signal Processing & STFT'), ('   4.2','Hearing Loss & Audiograms'),
    ('   4.3','Wavelet Denoising'), ('   4.4','MMSE-LSA Filtering'),
    ('   4.5','DNN Masking & CRM'), ('   4.6','FiLM Conditioning'),
    ('   4.7','Mamba: State Space Models'),
    ('5.','Proposed Methodology — 5-Model Comparison'),
    ('   5.1','System Overview'), ('   5.2','Model 1: Wavelet DWT'),
    ('   5.3','Model 2: MMSE-LSA'), ('   5.4','Model 3: 1D CNN'),
    ('   5.5','Model 4: U-Net + Attention + FiLM'), ('   5.6','Model 5: Mamba/SSM + FiLM'),
    ('   5.7','Training Strategy'),
    ('6.','Datasets'), ('7.','Evaluation Metrics'), ('8.','Project Timeline'),
    ('9.','Expected Results & Novel Contributions'),
    ('10.','Future Scope'), ('11.','Conclusion'), ('12.','References'),
]
for n,t in toc_items:
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    p.paragraph_format.left_indent=Inches(0.1 if n.startswith('   ') else 0)
    r=p.add_run(f'{n}    {t}'); r.font.size=Pt(10.5); r.font.name='Calibri'
    r.bold = not n.startswith('   ')
doc.add_page_break()

# ── 1. INTRODUCTION ────────────────────────────────────────────────────────
hd(doc,'1. Introduction'); add_hr(doc)
body(doc,'Hearing loss is one of the most prevalent sensory impairments worldwide. The WHO reports '
'approximately 466 million people — roughly 5% of the global population — have disabling hearing loss, '
'projected to reach 900 million by 2050. Hearing aids are the primary assistive technology.')
body(doc,'Despite advances in digital signal processing and deep learning, the most persistent complaint '
'among hearing aid users remains difficulty understanding speech in noisy environments. The "cocktail '
'party problem" is further complicated by damaged cochlear hair cells reducing frequency resolution.')
body(doc,'Modern hearing aids personalise amplification per audiogram but apply identical noise reduction '
'to every user. This is a fundamental mismatch — the frequencies most critical to intelligibility differ '
'across individuals. This project closes this gap via a 5-model comparison study culminating in two '
'personalised deep learning models that condition on the user\'s audiogram at inference time.')

# ── 2. PROBLEM STATEMENT ───────────────────────────────────────────────────
hd(doc,'2. Problem Statement'); add_hr(doc)
body(doc,'Given y(t) = s(t) + n(t) and a user audiogram a ∈ ℝ⁶ at {250,500,1000,2000,4000,8000} Hz, '
'learn a mapping:')
code(doc,'    f(Y, a) → Ŝ\n\n    Y = complex STFT of y(t), Ŝ ≈ complex STFT of s(t)\n'
'    via Complex Ratio Mask M_real + j*M_imag on both magnitude and phase.')
body(doc,'f must produce measurably different outputs for different audiograms. Primary validation: '
'ΔHASPI (personalized vs generic) > ΔSTOI (same comparison).')
hd(doc,'Key Constraints',2)
blt(doc,'Latency: < 10 ms per frame for real-time hearing aid deployment.')
blt(doc,'No retraining per user: audiogram is a runtime input, not a training-time choice.')
blt(doc,'Phase-awareness: CRM (Complex Ratio Mask) not magnitude-only IRM (Ideal Ratio Mask).')
blt(doc,'No clinical subjects: HASPI/HASQI are validated objective surrogates.')
doc.add_page_break()

# ── 3. LITERATURE REVIEW ───────────────────────────────────────────────────
hd(doc,'3. Literature Review & Research Gap'); add_hr(doc)
hd(doc,'3.1 Classical Speech Enhancement',2)
body(doc,'The MMSE-STSA estimator (Ephraim & Malah, 1985) is the most important classical algorithm — '
'running in most commercial hearing aids today. Its decision-directed SNR estimation suppresses noise '
'while reducing musical noise artefacts. Wavelet denoising (Donoho & Johnstone, 1994) adapts '
'frequency resolution, making it more suitable for non-stationary noise.')
hd(doc,'3.2 Deep Learning Approaches',2)
body(doc,'Conv-TasNet (Luo & Mesgarani, 2019) showed 1D waveform-domain processing achieves SOTA '
'performance with millisecond-scale latency. CMGAN (2022) and MetricGAN+ (2021) demonstrated adversarial '
'training with metric-predicting discriminators. SEMamba (Chao et al., 2024) achieved PESQ 3.69 on '
'VoiceBank-DEMAND — beating all Transformer variants — with linear O(T) complexity.')
hd(doc,'3.3 Hearing-Aid-Specific Research',2)
body(doc,'Diehl et al. (2023, Nature Scientific Reports) restored near-normal intelligibility for '
'hearing-impaired users using a generic U-Net — no audiogram conditioning. The Clarity Challenge '
'(University of Salford, 2021–2024) provides HA-specific datasets with real audiograms and HASPI/HASQI '
'evaluation. NeuroAMP (2025) conditions neural amplification on audiograms — for amplification only.')
box(doc,'Gap: No published system combines audiogram-conditioned noise reduction with HASPI-optimised '
'adversarial training. The Mamba+FiLM combination for hearing aids is also not published. '
'These are our two specific novel contributions.')
hd(doc,'3.4 Faculty Feedback Integration',2)
body(doc,'Following faculty review: (1) Wiener filter replaced by Wavelet DWT and MMSE-LSA, '
'(2) 1D CNN added as a waveform-domain DL baseline, (3) TIMIT and NOIZEUS added. '
'RASTA was researched and excluded — it is a feature-extraction technique for ASR, not for '
'producing listenable enhanced audio. Documented in literature review as a related technique.')
hd(doc,'3.5 Post-Research-Sweep Architecture Upgrades',2)
body(doc,'A 4-agent research sweep (August 2026) identified 6 architectural improvements:')
blt(doc,'CRM (Complex Ratio Mask): handles magnitude and phase — eliminates musical noise from IRM.')
blt(doc,'MetricGAN+ with HASPI discriminator: adversarial training optimising hearing-aid intelligibility.')
blt(doc,'ERB (Equivalent Rectangular Bandwidth) subband FiLM: 6→24 audiogram features aligned to auditory critical bands.')
blt(doc,'Mamba/SSM: replaces Transformer with O(T) linear model, PESQ 3.69 SOTA.')
blt(doc,'Self-Attention in bottleneck: global context without full Transformer cost.')
blt(doc,'Knowledge distillation: large model → small model for on-device deployment.')
doc.add_page_break()

# ── 4. BACKGROUND & THEORY ─────────────────────────────────────────────────
hd(doc,'4. Background & Theory'); add_hr(doc)
hd(doc,'4.1 Audio Signal Processing & STFT',2)
body(doc,'The STFT (Short-Time Fourier Transform) converts a 1D audio signal into a 2D complex '
'time-frequency spectrogram by applying the Discrete Fourier Transform to short overlapping windows:')
code(doc,'    X[k,n] = Σ x[m]·w[n-m]·e^(-j2πkm/N)\n'
'    k=freq bin, n=time frame, w=Hann window, N=512 → complex spec X ∈ ℂ^{257×T}\n'
'    Our models work with the full complex spectrum for CRM (phase-aware) masking.')
hd(doc,'4.2 Hearing Loss & Audiograms',2)
body(doc,'An audiogram maps hearing threshold (dB HL = decibels Hearing Level) against frequency. '
'0 dB HL = normal; 70 dB HL = must be 70 dB louder than normal to detect. '
'Model input: a = [a_250, a_500, a_1000, a_2000, a_4000, a_8000] / 100.0 → ∈ [0,1]^6')
tbl(doc,['Classification','Threshold (dB HL)','Impact'],[
    ['Normal','0–25','Can hear whispers'],
    ['Mild Loss','26–40','Misses quiet speech'],
    ['Moderate Loss','41–55','Misses conversational speech'],
    ['Moderately Severe','56–70','Requires amplification'],
    ['Severe / Profound','71–90+','Hears only loud sounds'],
])
hd(doc,'4.3 Wavelet Denoising',2)
body(doc,'DWT (Discrete Wavelet Transform) decomposes the signal into wavelet coefficients at multiple '
'time-scale levels. Noise = small coefficients everywhere; speech = large coefficients at specific scales. '
'Soft-thresholding removes noise while preserving speech:')
code(doc,'    coeffs = pywt.wavedec(signal, "db8", level=5)\n'
'    sigma = median(|coeffs[-1]|) / 0.6745    # noise std via MAD estimator\n'
'    threshold = sigma * sqrt(2*log(N))        # universal threshold\n'
'    clean = pywt.waverec(soft_threshold(coeffs, threshold), "db8")\n'
'    Library: PyWavelets  (pip install PyWavelets)')
hd(doc,'4.4 MMSE-LSA Filtering',2)
body(doc,'MMSE-LSA (Minimum Mean Square Error — Log Spectral Amplitude) uses decision-directed a priori '
'SNR estimation for an optimal perceptual gain. It is the algorithm inside commercial hearing aids today '
'(Oticon, Phonak, Signia). Reference: Ephraim & Malah (1985), IEEE TASLP.')
code(doc,'    xi(k,n) = alpha*A^2(k,n-1)/lambda_n + (1-alpha)*max(gamma-1,0)  # a priori SNR, alpha=0.98\n'
'    G(k,n)  = exp(0.5*E1(v)) / gamma                                   # MMSE-LSA gain\n'
'    v = xi*gamma/(1+xi),  gamma = instantaneous SNR')
hd(doc,'4.5 DNN Masking & CRM',2)
body(doc,'Deep learning predicts a mask per time-frequency bin. Our upgrade: Complex Ratio Mask (CRM) '
'predicts M_real and M_imag — modifying both magnitude AND phase:')
code(doc,'    # IRM (old) — magnitude only:\n'
'    IRM = |S|^2 / (|S|^2 + |N|^2)\n\n'
'    # CRM (ours) — phase-aware:\n'
'    S_real = M_real*Y_real - M_imag*Y_imag\n'
'    S_imag = M_real*Y_imag + M_imag*Y_real\n'
'    -> Correct phase -> eliminates musical noise artefacts')
hd(doc,'4.6 FiLM Conditioning',2)
body(doc,'FiLM (Feature-wise Linear Modulation, Perez et al. 2017) injects the audiogram into an '
'intermediate network layer via per-channel scale (gamma) and shift (beta) parameters:')
code(doc,'    FiLM Generator: a(6-dim) -> Linear(64) -> ReLU -> Linear(512)\n'
'                                                  -> gamma(256) + beta(256)\n\n'
'    FiLM Apply: h_out(b,c,h,w) = gamma(b,c) * h(b,c,h,w) + beta(b,c)\n'
'    Cost: ~131K extra parameters. Negligible.')
box(doc,'Verification: the model must produce different masks for different audiograms on the same '
'noisy clip. If masks are identical, FiLM is not working.')
hd(doc,'4.7 Mamba: State Space Models (SSM)',2)
body(doc,'Mamba (Gu & Dao, NeurIPS 2023) is a selective State Space Model with O(T) linear complexity — '
'unlike Transformers at O(T^2). Its state transition matrices are input-dependent, allowing selective '
'memory (like a gated RNN but much more powerful):')
code(doc,'    Standard SSM: h(t) = A*h(t-1) + B*x(t);  y(t) = C*h(t)   [fixed A,B,C]\n'
'    Mamba:        A(t),B(t),C(t) = f(x(t))           [input-dependent]\n'
'                  h(t) = A(t)*h(t-1) + B(t)*x(t)     [selective memory]\n'
'    SEMamba (2024): PESQ 3.69 on VoiceBank-DEMAND — current SOTA.\n'
'    Library: pip install mamba-ssm')
tbl(doc,['Property','Transformer','Mamba/SSM','Impact for Hearing Aids'],[
    ['Complexity','O(T^2)','O(T)','Mamba scales to continuous audio'],
    ['Real-time','Needs masking hacks','Naturally causal','Safe for real-time'],
    ['PESQ (SOTA)','~3.40','3.69','More accurate AND faster'],
    ['Parameters','10–30M','2–5M','Closer to on-device deployment'],
])
doc.add_page_break()

# ── 5. METHODOLOGY ─────────────────────────────────────────────────────────
hd(doc,'5. Proposed Methodology — 5-Model Comparison'); add_hr(doc)
hd(doc,'5.1 System Overview',2)
tbl(doc,['#','Model Name','Type','Domain','Audiogram?','Key Feature'],[
    ['1','Wavelet DWT Denoiser','Classical','Time-scale','No','Adaptive wavelet thresholding'],
    ['2','MMSE-LSA Filter','Classical','STFT','No','Industry standard, decision-directed SNR'],
    ['3','1D CNN (Conv-TasNet)','DL — generic','Waveform','No','Dilated convolutions, no STFT'],
    ['4★','U-Net + Attn + FiLM','DL — personalized','STFT/CRM','YES','HASPI-GAN + CRM + FiLM'],
    ['5','Mamba/SSM + FiLM','DL — personalized','STFT/CRM','YES','SOTA accuracy + linear O(T)'],
])
body(doc,'Classical -> Generic DL -> Personalized DL -> SOTA Personalized. '
'Models 4 and 5 must show DELTA_HASPI > DELTA_STOI vs Model 3 to validate audiogram conditioning.')
code(doc,'  Full Pipeline (Models 4 & 5):\n'
'  Noisy y(t)           Audiogram a [6-dim]\n'
'      |                      |\n'
'      v                      v\n'
'  [STFT]            [FiLM Generator: 6->64->512 -> gamma(256) + beta(256)]\n'
'      |                      |\n'
'      v                      |\n'
'  [Encoder CNN 4 stages] <---|\n'
'      |                      |\n'
'      v           <----------+\n'
'  [FiLM: h = gamma * h + beta]\n'
'      |\n'
'      v\n'
'  [Decoder CNN + skip connections]\n'
'      |\n'
'      v\n'
'  [CRM: M_real, M_imag] x [noisy complex spec]\n'
'      |\n'
'      v\n'
'  [ISTFT] -> Enhanced speech s_hat(t)')
hd(doc,'5.2 Model 1: Wavelet DWT Denoising',2)
blt(doc,'Algorithm: Daubechies-8 (db8) wavelet, 5-level decomposition, soft thresholding')
blt(doc,'Speed: <1 ms latency. No training required.')
blt(doc,'Advantage over Wiener filter: adaptive time-frequency resolution for non-stationary noise.')
blt(doc,'Library: PyWavelets (pywt)')
hd(doc,'5.3 Model 2: MMSE-LSA Filter',2)
blt(doc,'Algorithm: Ephraim & Malah (1985) decision-directed MMSE-LSA gain, alpha=0.98')
blt(doc,'Speed: <1 ms latency. No training required.')
blt(doc,'Significance: beating MMSE-LSA with DL is meaningful — it runs in commercial hearing aids today.')
hd(doc,'5.4 Model 3: 1D CNN — Conv-TasNet Style',2)
blt(doc,'Architecture: Encoder (1D Conv) -> TCN x8 (dilation 1,2,4..128) -> Mask -> Decoder')
blt(doc,'Input: raw waveform (B, T). No STFT — model learns its own frequency decomposition.')
blt(doc,'Loss: SI-SDR. No audiogram. Serves as best generic DL comparison point.')
blt(doc,'Library: Asteroid — ConvTasNet(n_src=1)')
hd(doc,'5.5 Model 4: U-Net + Self-Attention + FiLM',2)
blt(doc,'Encoder: 4 x 2D Conv stages stride=2, BatchNorm, LeakyReLU -> (B,256,F/16,T/16)')
blt(doc,'Bottleneck: Self-Attention (global context) + FiLM conditioning on audiogram')
blt(doc,'Decoder: 4 x ConvTranspose + skip connections -> (B, 2, F, T) = CRM output')
blt(doc,'Loss: SI-SDR + MetricGAN+(HASPI) discriminator loss (lambda annealed over 20 epochs)')
blt(doc,'FiLM: MLP (6->64->128->512) -> gamma, beta for 256 channels at bottleneck')
box(doc,'Novel: MetricGAN+ discriminator trained to predict HASPI scores (not PESQ). '
'Directly optimises the generator for hearing-impaired intelligibility. Not published.',BOX_TIP)
hd(doc,'5.6 Model 5: Mamba/SSM + FiLM',2)
blt(doc,'Architecture: STFT -> Conv extractor -> Mamba SSM blocks x4 -> FiLM -> CRM head -> ISTFT')
blt(doc,'Mamba blocks: selective SSM — input-dependent A, B, C matrices. O(T) complexity.')
blt(doc,'FiLM applied after Mamba blocks (same MLP as Model 4). Output: CRM — phase-aware masking.')
blt(doc,'Expected PESQ: ~3.6 (SEMamba: 3.69; ours adds FiLM). Latency: ~6 ms.')
blt(doc,'Library: mamba-ssm (pip install mamba-ssm)')
box(doc,'Novel: SEMamba (2024) PESQ 3.69 — no audiogram conditioning. '
'Adding FiLM to Mamba for hearing-aid personalization is our specific novel combination.',BOX_TIP)
hd(doc,'5.7 Training Strategy',2)
blt(doc,'Optimiser: Adam, lr=3e-4, cosine annealing; Batch: 16 x 4-second segments')
blt(doc,'Dynamic mixing: noise added on-the-fly at random SNR in [-5, +10] dB per step')
blt(doc,'Audiogram: synthetic profiles from WHO prevalence distribution — 9 types (mild/moderate/severe x flat/sloping/reverse)')
blt(doc,'Model 4 loss: SI-SDR + lambda x MetricGAN+(HASPI) discriminator; early stop on HASPI validation')
blt(doc,'Epochs: 100 with patience=10')
doc.add_page_break()

# ── 6. DATASETS ────────────────────────────────────────────────────────────
hd(doc,'6. Datasets'); add_hr(doc)
tbl(doc,['Dataset','Role','Size','Key Feature','Access'],[
    ['TIMIT','Clean speech training','6,300 utterances','Phoneme labels, 630 speakers','LDC / university library'],
    ['NOIZEUS','Primary evaluation','30 x 8 noises x 4 SNR','Purpose-built for SE evaluation','Free — ecs.utdallas.edu'],
    ['Clarity CEC2/3','Primary HA training','~11,000 scenes','Real listener audiograms included','Free — claritychallenge.org'],
    ['VoiceBank-DEMAND','Sanity check / pretraining','11,572 utterances','Standard SE benchmark','Free — SpeechBrain'],
    ['MUSAN + ESC-50','Noise augmentation','900+ clips','Diverse real-world noise types','Free — openslr.org'],
])
body(doc,'Three audiogram test profiles (held-out evaluation):')
tbl(doc,['Profile','Values (dB HL at 250–8000 Hz)','Description'],[
    ['A','[10,10,10,10,10,10]','Normal / near-normal hearing'],
    ['B','[10,15,20,45,70,85]','Severe high-frequency loss (most common)'],
    ['C','[60,60,60,60,60,60]','Flat severe loss'],
])
doc.add_page_break()

# ── 7. EVALUATION METRICS ──────────────────────────────────────────────────
hd(doc,'7. Evaluation Metrics'); add_hr(doc)
hd(doc,'Primary — Hearing-Aid Specific',2)
tbl(doc,['Metric','Full Name','Range','What It Measures'],[
    ['HASPI','Hearing Aid Speech Perception Index (v2)','0->1','Intelligibility for hearing-impaired listener. Takes audiogram. PRIMARY.'],
    ['HASQI','Hearing Aid Speech Quality Index (v2)','0->1','Quality/naturalness for hearing-impaired listener. Takes audiogram.'],
])
hd(doc,'Secondary — Standard Reference',2)
tbl(doc,['Metric','Full Name','Range','Role'],[
    ['STOI','Short-Time Objective Intelligibility','0->1','General intelligibility (no audiogram)'],
    ['PESQ','Perceptual Evaluation of Speech Quality','1->4.5','Cross-paper comparison'],
    ['SI-SDR','Scale-Invariant Signal-to-Distortion Ratio','dB higher=better','Training loss'],
    ['Latency','Inference time per 10 ms frame','ms lower=better','Real-time feasibility'],
    ['Params','Total trainable parameters','Count','Deployment size proxy'],
])
box(doc,'Key result: if DELTA_HASPI (Models 4,5 vs 3) > DELTA_STOI (same), audiogram conditioning '
'specifically benefits hearing-impaired listeners — not just audio quality in general. '
'This is the primary hypothesis of the project.')
doc.add_page_break()

# ── 8. TIMELINE ────────────────────────────────────────────────────────────
hd(doc,'8. Project Timeline'); add_hr(doc)
tbl(doc,['Phase','Task','Dates','Lead','Status'],[
    ['0','DSP Fundamentals — STFT, spectrograms, hearing loss simulation','Aug 15–22','Both','In Progress'],
    ['1','Audiology basics — audiogram generation, pyclarity cochlear model','Aug 23–29','Both','TODO'],
    ['2','Classical Baselines — Wavelet DWT + MMSE-LSA on NOIZEUS','Aug 30–Sep 6','Jwanil','TODO'],
    ['3','Data Pipeline — TIMIT + Clarity + NOIZEUS DataLoader, dynamic mixing','Sep 7–20','Namya','TODO'],
    ['4','1D CNN (Conv-TasNet) training and evaluation','Sep 21–Oct 4','Both','TODO'],
    ['5','U-Net + Attention + FiLM + CRM + MetricGAN+(HASPI) training','Oct 5–18','Both','TODO'],
    ['5b','Mamba/SSM + FiLM training and evaluation','Oct 19–25','Jwanil','TODO'],
    ['6','Full evaluation — all 5 models x all metrics x 3 audiogram profiles','Oct 26–28','Both','TODO'],
    ['7','Report writing, audio demos, final presentation','Oct 29–31','Both','TODO'],
    ['TARGET','SUBMIT — Deadline','November 1, 2026','—','DEADLINE'],
])
doc.add_page_break()

# ── 9. EXPECTED RESULTS ────────────────────────────────────────────────────
hd(doc,'9. Expected Results & Novel Contributions'); add_hr(doc)
hd(doc,'Expected Performance on NOIZEUS (Informed Estimates from Literature)',2)
tbl(doc,['Model','HASPI','PESQ','STOI','Latency','Params'],[
    ['Noisy (no processing)','~0.30','~1.6','~0.65','0 ms','—'],
    ['Model 1: Wavelet DWT','~0.45','~2.0','~0.74','<1 ms','—'],
    ['Model 2: MMSE-LSA','~0.55','~2.4','~0.78','<1 ms','—'],
    ['Model 3: 1D CNN','~0.65','~3.0','~0.86','~3 ms','~5M'],
    ['Model 4: U-Net+FiLM','~0.75','~3.3','~0.89','~8 ms','~15M'],
    ['Model 5: Mamba+FiLM','~0.80','~3.6','~0.91','~6 ms','~5M'],
])
body(doc,'* HASPI values for Profile B (severe HF loss). Estimates from literature. '
'Actual results reported honestly including negative findings.',italic=True)
hd(doc,'Novel Contributions',2)
num(doc,'Audiogram-conditioned noise reduction via FiLM. NeuroAMP (2025) conditions amplification only — not noise reduction. We are first to apply FiLM to the noise reduction stage.')
num(doc,'HASPI-optimised MetricGAN+ discriminator (strongest claim): training a GAN discriminator on HASPI rather than PESQ. Not published anywhere.')
num(doc,'Mamba/SSM + FiLM for hearing aids: SEMamba (2024) PESQ 3.69 with no audiogram conditioning. Our combination is novel.')
num(doc,'Systematic 5-model comparison with HA-specific metrics (HASPI/HASQI) across multiple audiogram profiles. Not published in this form.')
hd(doc,'Deliverables',2)
num(doc,'5 trained models with HASPI, HASQI, PESQ, STOI, SI-SDR, latency, parameter counts.')
num(doc,'Publication-quality results tables and plots.')
num(doc,'Audio demonstrations: before/after .wav for 3 noise types x 3 audiogram profiles.')
num(doc,'Full academic report (.docx + PDF).')
num(doc,'Reproducible open-source code at github.com/Jwanil/hearing-aid-speech-enhancement.')
doc.add_page_break()

# ── 10. FUTURE SCOPE ───────────────────────────────────────────────────────
hd(doc,'10. Future Scope'); add_hr(doc)
hd(doc,'Real Hardware Testing',2)
body(doc,'Deploy Mamba+FiLM on actual hearing aid DSP (Digital Signal Processor) chips — e.g., Arm '
'Cortex-M55 + Ethos-U55 NPU (Neural Processing Unit). Apply knowledge distillation and INT8 '
'quantization to fit hearing-aid-grade compute budgets. Validate <10 ms in silicon.')
hd(doc,'Research Paper Submission',2)
body(doc,'The HASPI-optimised MetricGAN+ discriminator is suitable for:')
blt(doc,'ICASSP 2027 — IEEE International Conference on Acoustics, Speech and Signal Processing (deadline ~Sep 2026)')
blt(doc,'Interspeech 2027 — primary international speech research conference')
blt(doc,'IEEE/ACM TASLP — Transactions on Audio, Speech, and Language Processing (journal)')
hd(doc,'Indian Language Testing',2)
body(doc,'Evaluate on Hindi, Gujarati, Bengali, Tamil using IndicSpeech or MUCS (Multilingual and '
'Code-switching ASR) datasets. Indian speech has different phoneme inventories and retroflex '
'consonants. Genuinely underexplored gap in the HA literature. Directly relevant to India\'s '
'63M+ hearing-impaired population.')
hd(doc,'Federated Learning',2)
body(doc,'Audiograms are private medical data. Federated Learning (FL) trains across many users without '
'centralising data — each device trains locally, shares only gradient updates. Compatible with FiLM: '
'each device updates only FiLM generator weights while sharing the backbone.')
hd(doc,'Companion Mobile Application',2)
body(doc,'A smartphone app where users photograph their audiogram. The app connects to the hearing aid via '
'Bluetooth LE (Low Energy), uploads the audiogram vector, and the model personalises itself instantly — '
'no clinic visit required.')
hd(doc,'Clinical User Study',2)
body(doc,'Recruit 20–30 hearing-impaired participants, play enhanced audio via calibrated headphones, '
'collect MOS (Mean Opinion Score) ratings. Compare vs commercial hearing aids. Gold standard validation '
'aligned with Clarity Challenge listener evaluation methodology.')
doc.add_page_break()

# ── 11. CONCLUSION ─────────────────────────────────────────────────────────
hd(doc,'11. Conclusion'); add_hr(doc)
body(doc,'This project proposes a targeted and technically rigorous solution to a well-defined gap: '
'personalised amplification is clinically standard in hearing aids; personalised noise reduction is not. '
'The 5-model comparison — classical baselines through SOTA deep learning — provides a complete and honest '
'quantitative picture of the state of the art and our contribution within it.')
body(doc,'The two specific novel contributions — (1) HASPI-optimised MetricGAN+ adversarial training, '
'and (2) Mamba/SSM with FiLM audiogram conditioning — represent genuine, testable advances not yet '
'published. Both are directly produceable from this minor project\'s deliverables, positioning this '
'work as a credible ICASSP 2027 submission candidate.')
body(doc,'If the personalized models achieve even a modest HASPI improvement over Model 3 — particularly '
'for severe audiogram profiles — it validates the core hypothesis and opens a practical path towards '
'truly individualised hearing aid processing for 466 million people worldwide.')
doc.add_page_break()

# ── 12. REFERENCES ─────────────────────────────────────────────────────────
hd(doc,'12. References'); add_hr(doc)
refs = [
    '[1]  Perez et al. (2018). FiLM: Visual Reasoning with a General Conditioning Layer. AAAI 2018.',
    '[2]  Diehl et al. (2023). Restoring speech intelligibility for hearing aid users with deep learning. Nature Scientific Reports, 13, 870.',
    '[3]  Chao et al. (2024). SEMamba: State Space Model for Speech Enhancement. arXiv:2405.06400.',
    '[4]  Gu & Dao (2023). Mamba: Linear-Time Sequence Modeling with Selective State Spaces. NeurIPS 2023.',
    '[5]  Luo & Mesgarani (2019). Conv-TasNet: Surpassing Ideal Time-Frequency Magnitude Masking. IEEE/ACM TASLP.',
    '[6]  Kates & Arehart (2021). HASPI v2. Speech Communication, 131, 35–46.',
    '[7]  Kates & Arehart (2014). HASQI v2. Journal of the Audio Engineering Society, 62(3).',
    '[8]  Ephraim & Malah (1985). Speech enhancement using MMSE log-spectral amplitude estimator. IEEE TASLP.',
    '[9]  Donoho & Johnstone (1994). Ideal spatial adaptation by wavelet shrinkage. Biometrika, 81(3).',
    '[10] Fu et al. (2021). MetricGAN+: An Improved Version of MetricGAN for Speech Enhancement. Interspeech.',
    '[11] Cao et al. (2022). CMGAN: Conformer-based Metric-GAN for Speech Enhancement. Interspeech.',
    '[12] Hu et al. (2020). DCCRN: Deep Complex Convolution Recurrent Network for Phase-Aware SE. Interspeech.',
    '[13] Clarity Challenge Team (2021–2024). The Clarity Enhancement Challenge. University of Salford / CVSSP.',
    '[14] NeuroAMP (2025). End-to-end Deep Neural Amplifier for Personalized Hearing Aids. IEEE TAI. arXiv:2502.10822.',
    '[15] Baby et al. (2021). CoNNear: Convolutional Neural-Network Model of Human Cochlear Mechanics. JASA-EL.',
    '[16] World Health Organization (2021). World report on hearing. WHO Press.',
    '[17] Loizou, P.C. (2007). Speech Enhancement: Theory and Practice. CRC Press.',
    '[18] Hermansky & Morgan (1994). RASTA processing of speech. IEEE TASLP, 2(4), 578–589.',
]
for ref in refs:
    p=doc.add_paragraph()
    p.paragraph_format.space_after=Pt(5); p.paragraph_format.left_indent=Inches(0.3)
    p.paragraph_format.first_line_indent=Inches(-0.3)
    r=p.add_run(ref); r.font.size=Pt(9.5); r.font.name='Calibri'

# ── SAVE ───────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                   'docs','Minor_Project_Proposal_Hearing_Aid_Speech_Enhancement.docx')
doc.save(out)
print(f'\n  Proposal saved to:\n    {out}\n')
