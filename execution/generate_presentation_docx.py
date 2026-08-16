"""
generate_presentation_docx.py
Generates a professionally formatted Word (.docx) document that mirrors
the content of docs/presentation.html — the project proposal for
"Audiogram-Personalised Speech Enhancement for Hearing Aids."

Run:  python execution/generate_presentation_docx.py
Output: docs/Project_Proposal_Hearing_Aid_Speech_Enhancement.docx

Requirements: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

# ─────────────────────────────────────────────────────────────────────────────
# COLOUR PALETTE (matches the HTML presentation)
# ─────────────────────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x07, 0x0d, 0x14)   # slide background (used for header shading)
DARK_BLUE  = RGBColor(0x0d, 0x1a, 0x2e)   # secondary surface
MID_BLUE   = RGBColor(0x00, 0x77, 0xb6)   # accent2
CYAN       = RGBColor(0x00, 0xb4, 0xd8)   # accent (primary)
TEAL       = RGBColor(0x00, 0xc4, 0xa0)   # teal
AMBER      = RGBColor(0xf4, 0xa2, 0x61)   # amber
GREEN      = RGBColor(0x52, 0xb7, 0x88)   # green
RED        = RGBColor(0xe6, 0x39, 0x46)   # red
WHITE      = RGBColor(0xff, 0xff, 0xff)
LIGHT_GRAY = RGBColor(0xe8, 0xf4, 0xf8)
DIM_TEXT   = RGBColor(0x7f, 0xa8, 0xbf)
SLIDE_BG   = RGBColor(0x0b, 0x16, 0x22)   # used for shaded section headers

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    """Set cell background colour via raw XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = str(rgb)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Add borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, color in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if color:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '12')
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=WHITE, size=10, italic=False):
    """Set cell text with formatting."""
    cell.text = ''
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size = Pt(size)
    run.font.name = 'Calibri'


def add_styled_paragraph(doc, text, level='body', color=None, bold=False,
                          italic=False, size=None, space_before=0, space_after=6,
                          align=WD_ALIGN_PARAGRAPH.LEFT):
    """Add a paragraph with full font control."""
    sizes = {'title': 26, 'h1': 20, 'h2': 16, 'h3': 13, 'eyebrow': 9, 'body': 10, 'small': 9}
    colors = {
        'title': WHITE, 'h1': CYAN, 'h2': CYAN, 'h3': LIGHT_GRAY,
        'eyebrow': CYAN, 'body': LIGHT_GRAY, 'small': DIM_TEXT
    }
    para = doc.add_paragraph()
    para.alignment = align
    pf = para.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold if bold is not None else (level in ('title', 'h1', 'h2', 'h3'))
    run.italic = italic
    run.font.color.rgb = color if color else colors.get(level, LIGHT_GRAY)
    run.font.size = Pt(size if size else sizes.get(level, 10))
    run.font.name = 'Calibri'
    return para


def add_slide_divider(doc, accent_color=CYAN):
    """Adds a coloured rule to simulate a slide header divider."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, accent_color)
    cell.text = ''
    cell.height = Cm(0.08)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_slide_header(doc, eyebrow: str, title: str, accent=CYAN):
    """Adds the standard eyebrow + title block for a slide."""
    add_slide_divider(doc, accent)
    # eyebrow
    p = add_styled_paragraph(doc, eyebrow.upper(), level='eyebrow',
                              color=accent, bold=True, space_before=4, space_after=2)
    p.runs[0].font.size = Pt(8.5)
    # title
    add_styled_paragraph(doc, title, level='h1', color=WHITE, bold=True,
                          size=18, space_before=0, space_after=10)


def add_bullet(doc, text, indent=0, color=LIGHT_GRAY, bold_prefix=None):
    """Add a bullet-style paragraph."""
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = para.add_run(bold_prefix + ': ')
        r.bold = True
        r.font.color.rgb = CYAN
        r.font.name = 'Calibri'
        r.font.size = Pt(10)
    r2 = para.add_run(text)
    r2.font.color.rgb = color
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)


def add_info_box(doc, label: str, content: str, accent=CYAN):
    """Add a shaded info box (accent left-border effect via table)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    # narrow accent strip
    accent_cell = table.cell(0, 0)
    set_cell_bg(accent_cell, accent)
    accent_cell.width = Cm(0.3)
    accent_cell.text = ''
    # content cell
    content_cell = table.cell(0, 1)
    set_cell_bg(content_cell, DARK_BLUE)
    content_cell.text = ''
    para = content_cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    if label:
        r1 = para.add_run(label + '\n')
        r1.bold = True
        r1.font.color.rgb = accent
        r1.font.size = Pt(8)
        r1.font.name = 'Calibri'
    r2 = para.add_run(content)
    r2.font.color.rgb = LIGHT_GRAY
    r2.font.size = Pt(9.5)
    r2.font.name = 'Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_code_block(doc, code: str):
    """Add a monospaced code-style block."""
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_bg(cell, RGBColor(0x05, 0x18, 0x28))
    set_cell_border(cell, top='00F5C4', bottom='00F5C4', left='00F5C4', right='00F5C4')
    cell.text = ''
    for line in code.strip().split('\n'):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(8.5)
        r.font.color.rgb = TEAL
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_page_break(doc):
    doc.add_page_break()


def styled_table(doc, headers, rows, col_widths=None):
    """Add a styled table with dark header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, MID_BLUE)
        cell_text(cell, h, bold=True, color=WHITE, size=9)

    # Data rows
    for ri, row_data in enumerate(rows):
        drow = table.rows[ri + 1]
        bg = RGBColor(0x0d, 0x1a, 0x2e) if ri % 2 == 0 else RGBColor(0x0f, 0x1f, 0x36)
        for ci, val in enumerate(row_data):
            cell = drow.cells[ci]
            set_cell_bg(cell, bg)
            is_first = ci == 0
            cell_text(cell, str(val), bold=is_first, color=WHITE if is_first else LIGHT_GRAY, size=9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table


def set_doc_background(doc):
    """Set document background colour."""
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), '070D14')
    doc.element.insert(0, bg)
    settings = doc.settings.element
    ds = OxmlElement('w:displayBackgroundShape')
    settings.append(ds)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN DOCUMENT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def build_document():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width  = Cm(29.7)   # A4 landscape
    section.page_height = Cm(21.0)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Dark background
    set_doc_background(doc)

    # Default Normal style
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.color.rgb = LIGHT_GRAY
    normal.font.size = Pt(10)

    # ── List Bullet style ──
    try:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Calibri'
        lb.font.size = Pt(10)
        lb.font.color.rgb = LIGHT_GRAY
    except Exception:
        pass

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    # Decorative header band
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, NAVY)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run('MINOR PROJECT PROPOSAL  ·  DEPARTMENT OF INFORMATION TECHNOLOGY  ·  2026')
    r.bold = True; r.font.color.rgb = DIM_TEXT; r.font.size = Pt(8); r.font.name = 'Calibri'
    doc.add_paragraph()

    add_styled_paragraph(doc, 'PROBLEM IN SIGNAL PROCESSING & DEEP LEARNING',
                         level='eyebrow', color=CYAN, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER, size=9, space_before=6, space_after=6)

    add_styled_paragraph(doc,
                         'Audiogram-Personalised Speech Enhancement\nfor Hearing Aids',
                         level='title', color=WHITE, bold=True, size=28,
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=10)

    add_styled_paragraph(doc,
                         'Using deep learning to personalise noise reduction based on each individual\'s unique hearing loss profile',
                         level='body', color=DIM_TEXT, size=11,
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=20)

    # Team table
    team_table = doc.add_table(rows=1, cols=2)
    team_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, name, roll in [
        (team_table.cell(0, 0), 'Jwanil Modi', '23BIT194'),
        (team_table.cell(0, 1), 'Namya Shah',  '23BIT027'),
    ]:
        set_cell_bg(cell, DARK_BLUE)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)
        r1 = p.add_run(name + '\n')
        r1.bold = True; r1.font.color.rgb = WHITE; r1.font.size = Pt(12); r1.font.name = 'Calibri'
        r2 = p.add_run(roll)
        r2.font.color.rgb = DIM_TEXT; r2.font.size = Pt(9); r2.font.name = 'Calibri'

    doc.add_paragraph()
    # Tag chips
    chip_text = '5-Model Comparison   |   HASPI / HASQI Evaluation   |   Mamba SOTA   |   November 1, 2026'
    add_styled_paragraph(doc, chip_text, color=DIM_TEXT, size=9,
                         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=4)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 2: THE PROBLEM
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'The Problem', 'Hearing Aids Amplify. They Don\'t Understand.')

    # Stats row
    stats = doc.add_table(rows=1, cols=2)
    stats.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, num, label in [
        (stats.cell(0, 0), '466 Million', 'people worldwide with disabling hearing loss (WHO)'),
        (stats.cell(0, 1), '#1 Complaint', 'Understanding speech in noisy environments'),
    ]:
        set_cell_bg(cell, DARK_BLUE)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(8)
        r1 = p.add_run(num + '\n'); r1.bold = True; r1.font.color.rgb = CYAN; r1.font.size = Pt(20); r1.font.name = 'Calibri'
        r2 = p.add_run(label); r2.font.color.rgb = DIM_TEXT; r2.font.size = Pt(9); r2.font.name = 'Calibri'
    doc.add_paragraph()

    add_info_box(doc, 'THE CORE PROBLEM',
                 'Modern hearing aids apply the same noise-reduction algorithm to every user — '
                 'regardless of their individual hearing loss profile.', CYAN)
    add_info_box(doc, 'WHY THIS IS WRONG',
                 'Every person\'s hearing loss is different. Person A loses high frequencies. Person B loses '
                 'low frequencies. The same algorithm cannot be optimal for both.', AMBER)
    add_info_box(doc, 'OUR SOLUTION',
                 'A neural network (a type of AI) that reads each person\'s audiogram (hearing test result) '
                 'and personalises its noise-reduction behaviour specifically for that user.', TEAL)

    add_styled_paragraph(doc, 'Three Different Audiogram Profiles', 'h3', bold=True, space_before=8)
    styled_table(doc,
        ['Frequency →', '250 Hz', '500 Hz', '1000 Hz', '2000 Hz', '4000 Hz', '8000 Hz'],
        [
            ['Normal',       '10 dB', '10 dB', '10 dB', '10 dB', '10 dB', '10 dB'],
            ['HF (High-Freq) Loss', '10 dB', '15 dB', '20 dB', '45 dB', '70 dB', '85 dB'],
            ['Flat Severe',  '60 dB', '60 dB', '60 dB', '60 dB', '60 dB', '60 dB'],
        ])
    add_styled_paragraph(doc, 'dB HL = decibels Hearing Level. Higher number = worse hearing at that frequency.',
                         'small', color=DIM_TEXT, space_before=0, space_after=4)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 3: WHAT IS AN AUDIOGRAM
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Key Concept', 'What Is an Audiogram?')

    add_styled_paragraph(doc,
        'An audiogram is the result of a clinical hearing test. It records the quietest sound '
        'a person can hear at 6 standard frequencies — from deep bass (250 Hz) to high treble (8000 Hz).',
        'body', space_before=0, space_after=8)

    add_styled_paragraph(doc, 'The 6 Test Frequencies', 'h3', bold=True, space_before=4)
    styled_table(doc,
        ['Frequency', 'Sound Region', 'Speech Content'],
        [
            ['250 Hz',  'Deep bass',       'Low vowel sounds, background hum'],
            ['500 Hz',  'Bass',            'Vowels (a, o, u)'],
            ['1000 Hz', 'Mid',             'Vowels, general speech energy'],
            ['2000 Hz', 'Mid-high',        'Speech clarity, consonant transitions'],
            ['4000 Hz', 'High',            'Consonants (s, f, sh, ch)'],
            ['8000 Hz', 'Very high',       'Fricatives (s, f), air sounds'],
        ])

    add_styled_paragraph(doc, 'Hearing Loss Classification', 'h3', bold=True, space_before=8)
    styled_table(doc,
        ['Range (dB HL)', 'Category', 'Practical Impact'],
        [
            ['0–25 dB HL',  'Normal',         'Can hear whispers'],
            ['25–40 dB HL', 'Mild',           'Misses quiet speech'],
            ['40–70 dB HL', 'Moderate',       'Misses normal conversational speech'],
            ['70+ dB HL',   'Severe/Profound','Only hears loud sounds'],
        ])

    add_code_block(doc, 'audiogram_vector = [10, 15, 20, 45, 70, 85]   # dB HL at 6 frequencies\n'
                        'normalised       = [0.10, 0.15, 0.20, 0.45, 0.70, 0.85]  # divided by 100 → model input')

    add_info_box(doc, 'WHY THIS IS OUR MODEL\'S INPUT',
                 'The audiogram tells us exactly which frequencies the person struggles with. '
                 'Hearing aids already personalise amplification per audiogram (using the NAL-R prescription formula). '
                 'We extend this to noise reduction — which is currently the same for every user.', TEAL)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 4: RESEARCH GAP
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Literature Review & Research Gap', 'What Exists — And What Doesn\'t')

    add_styled_paragraph(doc, 'What Already Exists', 'h3', bold=True, space_before=4)
    for txt, label in [
        ('Diehl et al., 2023 (Nature Scientific Reports) — Generic Deep Neural Network denoising '
         'that restored intelligibility to near-normal levels. No audiogram conditioning.', 'Paper 1'),
        ('Conv-TasNet (Luo & Mesgarani, 2019) — State-of-the-art waveform-domain separation. '
         'No personalisation.', 'Paper 2'),
        ('SEMamba (Chao et al., 2024) — Current SOTA (State of the Art): PESQ (Perceptual Evaluation '
         'of Speech Quality) score 3.69. No hearing-aid awareness.', 'Paper 3'),
        ('Clarity Challenge (University of Salford, 2021–2024) — Hearing-aid-specific benchmark. '
         'Winning systems use audiograms for scoring, not for conditioning the model.', 'Challenge'),
        ('NeuroAMP (Ahmed et al., 2025, IEEE TAI) — Audiogram conditioning for amplification, '
         'NOT for noise reduction.', 'Paper 4'),
        ('FiLM (Perez et al., 2017) — Feature-wise Linear Modulation conditioning mechanism. '
         'Originally for computer vision tasks.', 'Technique'),
    ]:
        add_bullet(doc, txt, bold_prefix=label)

    add_styled_paragraph(doc, 'The Gap We Target', 'h3', bold=True, space_before=10, color=TEAL)
    for label, content, color in [
        ('Gap 1',
         'No published system uses an audiogram as a conditioning input to a noise-reduction '
         'neural network. Audiograms are used for amplification or post-hoc evaluation only.',
         TEAL),
        ('Gap 2',
         'State-of-the-art architectures (SEMamba, CMGAN) are not evaluated on hearing-aid-specific '
         'metrics (HASPI/HASQI) and ignore the <10 ms latency constraint real hearing aids require.',
         CYAN),
        ('Our Specific Novelty',
         'Mamba/SSM + FiLM conditioning + MetricGAN+ trained with HASPI discriminator. '
         'No paper combines all three. The HASPI-GAN discriminator specifically is unpublished.',
         AMBER),
    ]:
        add_info_box(doc, label, content, color)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 5: 5-MODEL COMPARISON OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Our Approach', '5-Model Comparison Architecture')

    styled_table(doc,
        ['#', 'Model Name', 'Type', 'Key Feature', 'Audiogram?'],
        [
            ['1', 'Wavelet DWT (Discrete Wavelet Transform) Denoiser',
             'Classical', 'DWT soft-thresholding, non-stationary noise', 'No'],
            ['2', 'MMSE-LSA (Minimum Mean Square Error – Log Spectral Amplitude) Filter',
             'Classical', 'Industry standard in commercial hearing aids today', 'No'],
            ['3', '1D CNN (1-Dimensional Convolutional Neural Network) / Conv-TasNet',
             'Deep learning — generic', 'Waveform-domain, dilated convolutions', 'No'],
            ['4 ★', 'U-Net + Self-Attention + FiLM (Feature-wise Linear Modulation)',
             'Deep learning — personalised',
             'CRM (Complex Ratio Mask) + MetricGAN+ with HASPI discriminator', 'YES'],
            ['5', 'Mamba/SSM (State Space Model) + FiLM',
             'Deep learning — personalised',
             'SOTA accuracy (PESQ 3.69), linear complexity O(T)', 'YES'],
        ])

    add_info_box(doc, 'THE NARRATIVE',
                 'Classical → Generic deep learning → Personalised deep learning → SOTA personalised.\n'
                 'The key comparison: Models 4 & 5 vs Model 3 should show a larger HASPI improvement '
                 'than STOI (Short-Time Objective Intelligibility) improvement — proving that audiogram '
                 'conditioning specifically helps hearing-impaired listeners.', TEAL)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 6: MODELS 1 & 2 (CLASSICAL)
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Models 1 & 2', 'Classical Signal Processing Baselines', AMBER)

    add_styled_paragraph(doc, 'MODEL 1 — Wavelet DWT (Discrete Wavelet Transform) Denoising',
                         'h3', color=AMBER, bold=True, space_before=4)
    add_styled_paragraph(doc,
        'Converts the audio signal into a wavelet representation — similar to a spectrogram but with '
        'adaptive time-frequency resolution. Noise appears as small coefficients everywhere; speech '
        'concentrates in large coefficients at specific scales. We apply soft-thresholding to remove '
        'small coefficients (noise) while keeping large ones (speech).', 'body', space_after=6)
    add_code_block(doc,
        'DWT(signal) → wavelet coefficients at levels 1..L\n'
        'threshold λ = σ × √(2 × log(N))   [σ = noise std estimated via MAD]\n'
        'soft_threshold(coeff) → IDWT → clean signal\n'
        'Wavelet: db8 (Daubechies-8), level 5.  Library: PyWavelets (pywt)')

    add_styled_paragraph(doc, 'MODEL 2 — MMSE-LSA (Minimum Mean Square Error – Log Spectral Amplitude) Filter',
                         'h3', color=AMBER, bold=True, space_before=8)
    add_styled_paragraph(doc,
        'The algorithm actually running in commercial hearing aids today (Oticon, Signia, Phonak). '
        'Uses a "decision-directed" approach to estimate the signal-to-noise ratio frame-by-frame, '
        'then applies a perceptually-motivated gain that suppresses noise while preserving speech. '
        'No machine learning — pure signal processing. Virtually zero latency.', 'body', space_after=6)
    add_code_block(doc,
        'ξ(k,n) = α × Â²(k,n-1)/λ_n  +  (1-α) × max(γ-1, 0)\n'
        'α = 0.98,  ξ = a priori SNR (Signal-to-Noise Ratio),  γ = instantaneous SNR\n'
        'Reference: Ephraim & Malah (1985), IEEE TASLP')

    add_info_box(doc, 'WHY THESE BASELINES MATTER',
                 'These are not toy baselines. MMSE-LSA is the industry standard — it runs in virtually '
                 'every commercial hearing aid sold today. Beating it with deep learning is a meaningful result. '
                 '(Note: The Wiener filter was excluded as too basic — faculty feedback.)', AMBER)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 7: MODEL 3 — 1D CNN
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Model 3', '1D Convolutional Neural Network — Conv-TasNet Style')

    add_styled_paragraph(doc,
        'Works on the raw audio waveform directly — no STFT (Short-Time Fourier Transform) spectrogram needed. '
        'Learns its own internal frequency decomposition from data using 1D convolutions. '
        'Achieves very low latency. No audiogram input — serves as the "what does DL alone get us" baseline.',
        'body', space_after=8)

    add_styled_paragraph(doc, 'Architecture Block Diagram', 'h3', bold=True, space_before=4)
    styled_table(doc,
        ['Stage', 'Operation', 'Output Shape', 'Purpose'],
        [
            ['Encoder',  '1D Conv: kernel=16, stride=8, N=256 filters',
             '(B, 256, T/8)',   'Learned STFT replacement'],
            ['TCN Block 1', 'Dilated Conv, dilation=1',
             '(B, 256, T/8)',   'Local context, 1-frame window'],
            ['TCN Block 2–8', 'Dilated Conv, dilation=2,4..128',
             '(B, 256, T/8)',   'Growing receptive field up to ~32 ms'],
            ['Mask head', 'Linear → Sigmoid',
             '(B, 256, T/8)',   'Mask values in [0,1]'],
            ['Decoder', '1D ConvTranspose',
             '(B, 1, T)',       'Reconstruct enhanced waveform'],
        ])

    add_styled_paragraph(doc, 'Why Dilated Convolutions?', 'h3', bold=True, space_before=8)
    add_styled_paragraph(doc,
        'Standard 1D convolution with kernel=3 sees only 3 adjacent frames. To capture 1 second of '
        'context at 16 kHz you would need 500+ stacked layers. Dilation lets each layer skip frames, '
        'exponentially growing the receptive field with only 8 layers:',
        'body', space_after=4)
    add_code_block(doc,
        'dilation=1:   sees frames [t-1, t, t+1]      — 3 frames\n'
        'dilation=8:   sees frames [t-8, t, t+8]      — 17 frames\n'
        'dilation=128: sees frames [t-128, t, t+128]  — 257 frames\n'
        '→ Total receptive field across 8 blocks: 511 frames ≈ 32 ms at 16 kHz')

    add_info_box(doc, 'ROLE IN OUR COMPARISON',
                 'Model 3 answers: "What does a state-of-the-art generic DNN achieve without any personalisation?" '
                 'Models 4 and 5 must beat this using the audiogram. '
                 'The improvement (HASPI gain) = the measured benefit of personalisation.',
                 CYAN)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 8: MODEL 4 — U-Net + FiLM
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Model 4 — Core Contribution', 'U-Net + Self-Attention + FiLM Conditioning', TEAL)

    add_styled_paragraph(doc, 'Processing Pipeline', 'h3', bold=True, space_before=4)
    styled_table(doc,
        ['Step', 'Operation', 'Input → Output'],
        [
            ['1', 'STFT (Short-Time Fourier Transform)',
             'Waveform (B, T) → Complex Spectrogram (B, F=257, T_frames)'],
            ['2', 'Encoder — 4 × 2D CNN stages + stride=2',
             '(B, 1, F, T) → (B, 256, F/16, T/16)'],
            ['3', 'Self-Attention in bottleneck',
             '(B, 256, H, W) → same shape, global context added'],
            ['4', 'FiLM Layer ← Audiogram vector [a₁..a₆]',
             'γ × features + β  (personalisation applied here)'],
            ['5', 'Decoder — 4 × ConvTranspose + skip connections',
             '(B, 256, H/16, W/16) → (B, 2, F, T)'],
            ['6', 'CRM (Complex Ratio Mask) application',
             'M_real, M_imag × Noisy complex spectrum'],
            ['7', 'ISTFT (Inverse Short-Time Fourier Transform)',
             'Enhanced complex spectrum → Enhanced waveform'],
        ])

    add_styled_paragraph(doc, 'Three Stacked Innovations', 'h3', bold=True, space_before=8)
    for label, content, color in [
        ('INNOVATION 1: FiLM Conditioning (Personalisation)',
         'A small auxiliary MLP (Multi-Layer Perceptron: 6→64→128→512) maps the audiogram vector to '
         'scale (γ) and shift (β) parameters. Applied to bottleneck features: output = γ × features + β. '
         'Different audiograms produce different γ and β → different internal processing for each user.', TEAL),
        ('INNOVATION 2: Complex Ratio Mask (CRM)',
         'Predicts TWO mask channels — M_real and M_imag — and multiplies them with the complex spectrogram. '
         'This modifies both magnitude AND phase simultaneously, eliminating "musical noise" artifacts '
         'caused by the standard IRM (Ideal Ratio Mask) that only touches magnitude.', CYAN),
        ('INNOVATION 3: MetricGAN+ with HASPI Discriminator ★ (Strongest Novelty)',
         'A discriminator network learns to predict HASPI (Hearing Aid Speech Perception Index) scores. '
         'The generator (our U-Net) trains adversarially to maximise predicted HASPI — directly optimising '
         'for hearing-impaired listener intelligibility. Standard MetricGAN+ uses PESQ; using HASPI is '
         'our specific unpublished contribution.', AMBER),
    ]:
        add_info_box(doc, label, content, color)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 9: FiLM EXPLAINED
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'FiLM — Feature-wise Linear Modulation', 'How the Model Reads the Audiogram')

    add_styled_paragraph(doc,
        'FiLM is a conditioning mechanism that allows a neural network to modify its own internal '
        'behaviour based on a side input — in our case, the audiogram. Rather than concatenating the '
        'audiogram to the input (which doesn\'t work well for 2D feature maps), FiLM generates '
        'per-channel scale-and-shift parameters that directly modulate the feature maps at the bottleneck.',
        'body', space_after=8)

    add_styled_paragraph(doc, 'FiLM Operation', 'h3', bold=True, space_before=4)
    add_code_block(doc,
        'audiogram a = [a₁, a₂, a₃, a₄, a₅, a₆]   ← 6 normalised hearing thresholds\n'
        '\n'
        'FiLM Generator (small MLP):\n'
        '   a (6-dim) → Linear(64) → ReLU → Linear(128) → ReLU → Linear(512)\n'
        '                                                          → γ (256) + β (256)\n'
        '\n'
        'FiLM Application:\n'
        '   output(b, c, h, w) = γ(b, c) × features(b, c, h, w) + β(b, c)\n'
        '   (γ and β broadcast over spatial dimensions h and w)')

    add_styled_paragraph(doc, 'What FiLM Learns in Practice', 'h3', bold=True, space_before=8)
    styled_table(doc,
        ['Audiogram Type', 'Expected FiLM Behaviour', 'Effect'],
        [
            ['Severe High-Freq Loss\n[10,15,20,45,70,85]',
             'γ/β amplify high-freq noise-detection channels',
             'More aggressive HF noise removal — person can\'t hear HF anyway'],
            ['Flat Moderate Loss\n[40,40,40,40,40,40]',
             'γ/β adapt uniformly across all channels',
             'Balanced enhancement across all frequencies'],
            ['Near-Normal\n[10,10,10,10,10,10]',
             'γ ≈ 1, β ≈ 0 (near-identity)',
             'Minimal modification — person barely needs personalisation'],
        ])

    add_styled_paragraph(doc, 'Critical Verification Test', 'h3', bold=True, space_before=8, color=AMBER)
    add_code_block(doc,
        '# Run the same noisy clip with two different audiograms\n'
        'mask_normal = model(noisy, audiogram=[10,10,10,10,10,10])\n'
        'mask_hf     = model(noisy, audiogram=[10,15,20,45,70,85])\n'
        '\n'
        '# MUST be numerically different — if identical, FiLM is not being used\n'
        'assert torch.max(torch.abs(mask_normal - mask_hf)) > 0.01\n'
        'print("✅ Personalisation verified")')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 10: MODEL 5 — MAMBA
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Model 5 — State-of-the-Art', 'Mamba / State Space Model (SSM) + FiLM', GREEN)

    add_styled_paragraph(doc, 'Why Mamba Beats Transformers for Audio Processing', 'h3',
                         bold=True, color=GREEN, space_before=4)
    styled_table(doc,
        ['Property', 'Standard Transformer', 'Mamba / SSM', 'Why It Matters for Hearing Aids'],
        [
            ['Complexity',       'O(T²) — quadratic',  'O(T) — linear',
             'Hearing aids process continuous audio; quadratic cost explodes with time'],
            ['Real-time mode',   'Requires future masking hack', 'Naturally causal (left→right)',
             'Real-time = cannot look at future frames'],
            ['Best PESQ score',  '~3.40',              '3.69 (SEMamba, 2024)',
             'More accurate AND faster — not a tradeoff'],
            ['Parameter count',  '10–30M typical',     '2–5M',
             'Smaller models fit on hearing aid hardware'],
        ])

    add_styled_paragraph(doc, 'What Is a State Space Model (SSM)?', 'h3', bold=True, space_before=8)
    add_styled_paragraph(doc,
        'A State Space Model describes a system with memory: the current output depends on '
        'everything that happened before, compressed into a hidden state h(t).', 'body', space_after=4)
    add_code_block(doc,
        'Standard SSM (fixed matrices):\n'
        '   h(t) = A × h(t-1) + B × x(t)       ← hidden state update\n'
        '   y(t) = C × h(t)                     ← output\n'
        '\n'
        'Mamba (selective — matrices are input-dependent):\n'
        '   A(t), B(t), C(t) = f(x(t))          ← computed from current input!\n'
        '   h(t) = A(t) × h(t-1) + B(t) × x(t) ← selective memory\n'
        '   y(t) = C(t) × h(t)                  ← selective output\n'
        '\n'
        '→ When it hears a speech onset: opens gate, strongly updates h\n'
        '→ When it hears steady noise:   closes gate, barely updates h')

    add_info_box(doc, 'SIMPLE ANALOGY',
                 'An old RNN (Recurrent Neural Network) is like a conveyor belt that moves everything '
                 'forward at the same speed. Mamba is like a smart gatekeeper who decides what gets '
                 'remembered and what gets discarded — based on what is currently being processed.', GREEN)

    add_info_box(doc, 'OUR ADDITION: FiLM CONDITIONING',
                 'SEMamba (2024) has no audiogram input. We add a FiLM layer after the Mamba blocks. '
                 'This personalises the SSM\'s processing per audiogram. The Mamba+FiLM combination '
                 'for hearing aids is not published. Reference: Chao et al., 2024 (SEMamba).', TEAL)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 11: FULL SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'System Architecture', 'Complete Pipeline — Block by Block')

    add_styled_paragraph(doc,
        'The system has two inputs: (1) noisy speech waveform from the hearing aid microphone, '
        'and (2) the user\'s audiogram vector. The output is enhanced speech personalised for that user.',
        'body', space_after=8)

    styled_table(doc,
        ['Block', 'Operation', 'Input Shape', 'Output Shape', 'Notes'],
        [
            ['INPUT A', 'Noisy waveform from mic',
             '(B, T) samples', '(B, T)', '16 kHz, 10ms frames'],
            ['INPUT B', 'Audiogram vector',
             '(B, 6) dB HL values', '(B, 6) normalised', 'Normalised to [0, 1]'],
            ['STFT', 'Short-Time Fourier Transform',
             '(B, T)', '(B, F=257, T_frames) complex', 'Window=512, hop=128'],
            ['Encoder', '4× 2D Conv stages + BatchNorm + LeakyReLU',
             '(B, 1, 257, T)', '(B, 256, 16, T/16)', 'Each stage: stride=2'],
            ['Self-Attention', 'Q/K/V attention over flattened spatial dims',
             '(B, 256, H, W)', '(B, 256, H, W)', 'Global context, gamma=0 init'],
            ['FiLM', 'Audiogram MLP → γ, β → scale+shift features',
             '(B, 256, H, W) + (B,6)', '(B, 256, H, W)', 'Personalisation step'],
            ['Decoder', '4× ConvTranspose + skip connections from encoder',
             '(B, 256, H, W)', '(B, 2, 257, T)', '2 channels = M_real, M_imag'],
            ['CRM Apply', '(M_real + j×M_imag) × (Y_real + j×Y_imag)',
             '(B, 2, F, T) + noisy spec', '(B, F, T) complex', 'Phase-aware masking'],
            ['ISTFT', 'Inverse Short-Time Fourier Transform',
             '(B, F, T) complex', '(B, T) waveform', 'Enhanced speech'],
        ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 12: DATASETS
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Data', 'Datasets and Data Strategy')

    styled_table(doc,
        ['Dataset', 'Role', 'Size', 'Key Feature'],
        [
            ['TIMIT', 'Clean speech — training', '6,300 utterances', 'Phoneme labels, balanced speakers'],
            ['NOIZEUS', 'Evaluation benchmark', '30 sentences × 8 noises × 4 SNRs',
             'Standard SE (Speech Enhancement) evaluation corpus'],
            ['Clarity CEC2/3', 'Primary HA (Hearing Aid) training', '~11,000 scenes',
             'Real listener audiograms included, purpose-built for HA research'],
            ['VoiceBank-DEMAND', 'Secondary / sanity check', '11,572 utterances',
             'Standard SE benchmark — for cross-paper comparison'],
            ['MUSAN + ESC-50', 'Noise augmentation', '900+ noise clips',
             'Diverse real-world noise types for dynamic mixing'],
        ])

    add_styled_paragraph(doc, 'Dynamic Mixing Strategy (Training)', 'h3', bold=True, space_before=8, color=CYAN)
    add_code_block(doc,
        '# Instead of pre-creating fixed noisy files, mix on-the-fly every epoch:\n'
        'SNR    = random.uniform(-5, +10)   # dB, drawn fresh each step\n'
        'noise  = random.choice(MUSAN)       # random noise file\n'
        'noisy  = clean + 10^(-SNR/20) × noise\n'
        '# → Effective infinite training data. Model never memorises a fixed noise file.')

    add_info_box(doc, 'AUDIOGRAM STRATEGY FOR TIMIT-BASED TRAINING',
                 'TIMIT has no audiograms. We sample synthetic audiograms from a realistic statistical '
                 'distribution based on WHO/clinical hearing loss prevalence data. Clarity Challenge '
                 'provides real listener audiograms for training on HA-specific data.', AMBER)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 13: EVALUATION METRICS
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Evaluation', 'Metrics — Why Standard Metrics Aren\'t Enough')

    add_styled_paragraph(doc, 'Hearing-Aid Specific Metrics (Our Primary Metrics)', 'h3',
                         bold=True, color=TEAL, space_before=4)
    styled_table(doc,
        ['Metric', 'Full Name', 'Takes Audiogram?', 'Range', 'What It Measures'],
        [
            ['HASPI', 'Hearing Aid Speech Perception Index', 'YES ★', '0→1',
             'Predicted intelligibility for hearing-impaired listener'],
            ['HASQI', 'Hearing Aid Speech Quality Index', 'YES ★', '0→1',
             'Predicted quality (naturalness) for hearing-impaired listener'],
        ])

    add_styled_paragraph(doc, 'Standard Reference Metrics', 'h3', bold=True, space_before=8, color=CYAN)
    styled_table(doc,
        ['Metric', 'Full Name', 'Audiogram?', 'Range', 'Purpose'],
        [
            ['STOI',   'Short-Time Objective Intelligibility', 'No', '0→1',
             'Intelligibility for normal-hearing listeners'],
            ['SI-SDR', 'Scale-Invariant Signal-to-Distortion Ratio', 'No', 'dB (higher = better)',
             'Signal quality, used as training loss'],
            ['PESQ',   'Perceptual Evaluation of Speech Quality', 'No', '1→4.5',
             'Quality metric — used in SE literature for cross-paper comparison'],
            ['Latency','Processing delay per 10 ms frame', 'No', 'ms (lower = better)',
             'Must be <10 ms for real-time hearing aid deployment'],
        ])

    add_info_box(doc, 'THE KEY SCIENTIFIC RESULT WE EXPECT',
                 'For Models 4 and 5 vs Model 3 (generic DNN):\n\n'
                 '   ΔHASPI  >>  ΔSTOI\n\n'
                 'Meaning: the personalised model helps hearing-impaired listeners MORE than it helps '
                 'general listeners. This proves audiogram conditioning provides hearing-specific benefit, '
                 'not just generic noise removal.', TEAL)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 14: EXPECTED RESULTS
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Expected Outcomes', 'Results We Anticipate')

    add_styled_paragraph(doc, 'Expected Performance on NOIZEUS Benchmark',
                         'h3', bold=True, space_before=4)
    styled_table(doc,
        ['Model', 'HASPI ↑', 'PESQ ↑', 'STOI ↑', 'Latency ↓', 'Params'],
        [
            ['Noisy (no processing)', '~0.30', '~1.6', '~0.65', '0 ms', '—'],
            ['Model 1: Wavelet DWT',  '~0.45', '~2.0', '~0.74', '<1 ms', '—'],
            ['Model 2: MMSE-LSA',     '~0.55', '~2.4', '~0.78', '<1 ms', '—'],
            ['Model 3: 1D CNN',       '~0.65', '~3.0', '~0.86', '~3 ms', '~5M'],
            ['Model 4: U-Net+FiLM ★','~0.75', '~3.3', '~0.89', '~8 ms', '~15M'],
            ['Model 5: Mamba+FiLM',  '~0.80', '~3.6', '~0.91', '~6 ms', '~5M'],
        ])

    add_styled_paragraph(doc, '* HASPI values are for the severe HF loss audiogram [10,15,20,45,70,85]. '
                         'Numbers are informed estimates from published literature, not guaranteed results.',
                         'small', color=DIM_TEXT, space_before=2, space_after=10)

    add_styled_paragraph(doc, 'Three Test Audiogram Profiles', 'h3', bold=True, space_before=8)
    styled_table(doc,
        ['Profile ID', 'Values (dB HL at 250–8000 Hz)', 'Description', 'Expected personalisation gain'],
        [
            ['Profile A', '[10, 10, 10, 10, 10, 10]', 'Normal / near-normal hearing',
             'Smallest — model barely needs to adapt'],
            ['Profile B', '[10, 15, 20, 45, 70, 85]', 'Severe high-frequency loss (most common)',
             'Largest — strong HF suppression learned'],
            ['Profile C', '[60, 60, 60, 60, 60, 60]', 'Flat severe loss',
             'Moderate — uniform adaptation across frequencies'],
        ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 15: TIMELINE
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Project Timeline', 'August 2026 → November 1, 2026 (77 Days)')

    styled_table(doc,
        ['Phase', 'Name', 'Dates', 'Lead', 'Status'],
        [
            ['0',  'DSP (Digital Signal Processing) Fundamentals',      'Aug 15–22',     'Both', '🟡 In Progress'],
            ['1',  'Audiology + Audiogram Generation',                   'Aug 23–29',     'Both', '⬜ TODO'],
            ['2',  'Classical Baselines: Wavelet DWT + MMSE-LSA',        'Aug 30–Sep 6',  'Jwanil', '⬜ TODO'],
            ['3',  'Data Pipeline (TIMIT + Clarity + NOIZEUS)',          'Sep 7–20',      'Namya', '⬜ TODO'],
            ['4',  '1D CNN Model (Conv-TasNet)',                         'Sep 21–Oct 4',  'Both', '⬜ TODO'],
            ['5',  'U-Net + Attention + FiLM (Core Contribution)',       'Oct 5–18',      'Both', '⬜ TODO'],
            ['5b', 'Mamba/SSM + FiLM (SOTA Model)',                     'Oct 19–25',     'Jwanil', '⬜ TODO'],
            ['6',  'Full Evaluation — all 5 models, all metrics',        'Oct 26–28',     'Both', '⬜ TODO'],
            ['7',  'Report + Audio Demo',                                'Oct 29–31',     'Both', '⬜ TODO'],
            ['🎯', 'SUBMIT', 'November 1, 2026', '—', '🎯 DEADLINE'],
        ])

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 16: TOOLS & LIBRARIES
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Implementation', 'Tools, Libraries & Software Stack')

    styled_table(doc,
        ['Library', 'Version', 'Purpose', 'Install'],
        [
            ['PyTorch',      '≥2.1',  'Core deep learning framework',           'pip install torch torchaudio'],
            ['Asteroid',     '≥0.6',  'Conv-TasNet, SI-SDR loss, source sep tools', 'pip install asteroid'],
            ['SpeechBrain',  '≥1.0',  'MetricGAN+ reference implementation',   'pip install speechbrain'],
            ['mamba-ssm',    '≥1.0',  'Mamba State Space Model layers',         'pip install mamba-ssm'],
            ['pyclarity',    '≥0.4',  'HASPI & HASQI metrics, HA simulation',  'pip install pyclarity'],
            ['pystoi',       '≥0.4',  'STOI metric',                            'pip install pystoi'],
            ['pesq',         '≥0.0.4','PESQ metric',                            'pip install pesq'],
            ['PyWavelets',   '≥1.4',  'DWT denoising (Model 1)',                'pip install PyWavelets'],
            ['librosa',      '≥0.10', 'Audio analysis & visualisation',         'pip install librosa'],
            ['TensorBoard',  '≥2.14', 'Training loss & metric monitoring',      'pip install tensorboard'],
        ])

    add_code_block(doc,
        'pip install torch torchaudio asteroid speechbrain pyclarity pystoi '
        'pesq PyWavelets mamba-ssm librosa matplotlib')

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 17: NOVELTY
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Novelty & Contribution', 'What Makes This Project Original')

    for num, label, content, color in [
        ('1', 'Audiogram-Conditioned Noise Reduction',
         'NeuroAMP (2025) uses audiograms for amplification. No published system uses audiogram '
         'conditioning for the noise reduction step. We are the first to apply FiLM to hearing-aid '
         'noise reduction specifically.', TEAL),
        ('2 ★', 'HASPI-Optimised GAN Training — Strongest Claim',
         'MetricGAN+ trains a discriminator on PESQ (Perceptual Evaluation of Speech Quality). '
         'We train our discriminator on HASPI (Hearing Aid Speech Perception Index) — which takes '
         'the audiogram as input and measures intelligibility for hearing-impaired listeners. '
         'This specific combination has not been published anywhere.', AMBER),
        ('3', 'Mamba/SSM + FiLM for Hearing Aids',
         'SEMamba (Chao et al., 2024) achieves SOTA PESQ 3.69 but has no audiogram conditioning. '
         'We add FiLM to a SEMamba-style architecture for personalised hearing aid enhancement. '
         'The combination is novel.', GREEN),
        ('4', 'Systematic 5-Model Comparison with HA-Specific Metrics',
         'A rigorous comparison from classical DSP through SOTA deep learning, evaluated on '
         'HASPI/HASQI across multiple audiogram profiles, has not been published in this form. '
         'Directly useful to the field regardless of the primary novel results.', CYAN),
    ]:
        add_info_box(doc, f'CLAIM {num}: {label}', content, color)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 18: FUTURE SCOPE
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Future Scope', 'Where This Goes Next')

    for icon, label, content, color in [
        ('🔬', 'Real Hardware Testing',
         'Deploy the Mamba+FiLM model on actual hearing aid DSP (Digital Signal Processor) chips — '
         'e.g. Arm Cortex-M55 + Ethos-U55 NPU (Neural Processing Unit). Benchmark true on-device '
         'latency, power draw, and memory footprint. Validate that the <10ms constraint holds in '
         'silicon, not just simulation.', TEAL),
        ('📄', 'Research Paper Submission',
         'The HASPI-optimised MetricGAN+ discriminator + audiogram FiLM conditioning is an original '
         'contribution. Target venues: ICASSP 2027 (IEEE International Conference on Acoustics, Speech '
         'and Signal Processing), Interspeech 2027, or IEEE/ACM TASLP (Transactions on Audio, Speech, '
         'and Language Processing). Deadline for ICASSP 2027: ~September 2026.', CYAN),
        ('🌐', 'Indian Language Testing',
         'Evaluate robustness on Indian languages — Hindi, Gujarati, Bengali, Tamil — using '
         'IndicSpeech or MUCS (Multilingual and Code-switching ASR) datasets. Indian speech has '
         'different phoneme inventories and retroflex consonants. This is a genuinely underexplored '
         'gap in the hearing-aid literature, and directly relevant to India\'s 63M+ hearing-impaired population.', AMBER),
        ('🔐', 'Federated Learning',
         'Audiograms are private medical data. Federated Learning (FL) trains the model across many '
         'hearing aid users without sending raw audio or audiograms to a central server — each device '
         'trains locally and shares only gradient updates. Privacy-preserving personalisation at scale.', GREEN),
        ('📱', 'Companion Mobile App',
         'A smartphone app where users input or import their audiogram (photo of their audiogram chart). '
         'The app connects to the hearing aid via Bluetooth LE, uploads the audiogram vector, and the '
         'model re-conditions itself in real time. Audiogram update = instant personalisation with no '
         'clinic visit required.', RGBColor(0xc7, 0x7d, 0xff)),
        ('🏥', 'Clinical User Study',
         'Recruit 20–30 hearing-impaired participants, play enhanced audio via calibrated headphones, '
         'collect subjective intelligibility and quality ratings. Compare our system vs commercial '
         'hearing aids. This is the gold standard validation — aligns with the Clarity Challenge\'s '
         'listener evaluation methodology (MOS: Mean Opinion Score).', RGBColor(0xf7, 0x25, 0x85)),
    ]:
        add_info_box(doc, f'{icon}  {label}', content, color)

    add_info_box(doc, '★ THE HIGHEST-IMPACT NEXT STEP',
                 'Submit a paper to ICASSP 2027 on the HASPI-optimised MetricGAN+ finding. '
                 'This is the specific result that has not been published — and it is directly testable '
                 'from the deliverables of this minor project.', TEAL)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 19: CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'Conclusion', 'Summary & Why This Project Matters')

    add_info_box(doc, 'THE PROBLEM IN ONE LINE',
                 '466 million hearing aid users get generic noise reduction. '
                 'We personalise it per audiogram.', CYAN)
    add_info_box(doc, 'OUR APPROACH',
                 '5-model comparison from classical baselines to SOTA deep learning. FiLM conditioning '
                 'makes the model "read" each person\'s audiogram and adapt its noise-removal behaviour '
                 'specifically for them.', TEAL)
    add_info_box(doc, 'THE SCIENTIFIC CLAIM',
                 'HASPI improvement from audiogram conditioning will be larger than STOI improvement — '
                 'proving personalisation specifically benefits hearing-impaired listeners, not just '
                 'audio quality in general.', GREEN)

    add_styled_paragraph(doc, 'Deliverables by November 1, 2026', 'h3', bold=True, space_before=8)
    for item in [
        '5 trained and evaluated models with documented metrics (HASPI, HASQI, STOI, SI-SDR, PESQ, latency)',
        'Results table and 4 publication-quality plots (bar charts, scatter plots)',
        'Audio demonstrations: before/after .wav files for 3 noise types × 3 audiogram profiles',
        'Full written academic report (.docx and PDF)',
        'Reproducible open-source code on GitHub (github.com/Jwanil)',
        'Possibility of ICASSP 2027 submission for the HASPI-GAN finding',
    ]:
        add_bullet(doc, item)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 20: REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    add_slide_header(doc, 'References', 'Key Literature')

    refs = [
        ['1',  'Diehl et al.',          'Nature Scientific Reports',   '2023',
         '"Restoring speech intelligibility for hearing aid users with deep learning"'],
        ['2',  'Perez et al.',           'AAAI',                        '2018',
         '"FiLM: Visual Reasoning with a General Conditioning Layer"'],
        ['3',  'Chao et al.',            'arXiv',                       '2024',
         '"SEMamba: State Space Model for Speech Enhancement" — PESQ 3.69 SOTA'],
        ['4',  'Kates & Arehart',        'JASA',                        '2021',
         '"HASPI v2 / HASQI v2 — hearing aid intelligibility/quality indices"'],
        ['5',  'Ahmed et al.',           'IEEE TAI',                    '2025',
         '"NeuroAMP: End-to-end Deep Neural Amplifier for Personalized Hearing Aids"'],
        ['6',  'Luo & Mesgarani',        'IEEE/ACM TASLP',              '2019',
         '"Conv-TasNet: Surpassing Ideal Time-Frequency Masking for Speech Separation"'],
        ['7',  'Cao et al.',             'Interspeech',                 '2022',
         '"CMGAN: Conformer-based MetricGAN for Speech Enhancement"'],
        ['8',  'Fu et al.',              'Interspeech',                 '2021',
         '"MetricGAN+: An Improved Version of MetricGAN for Speech Enhancement"'],
        ['9',  'Hu et al.',              'Interspeech',                 '2020',
         '"DCCRN: Deep Complex Convolution Recurrent Network for Phase-Aware SE"'],
        ['10', 'Gu & Dao',              'NeurIPS',                     '2023',
         '"Mamba: Linear-Time Sequence Modeling with Selective State Spaces"'],
        ['11', 'Clarity Challenge',      'Univ. Salford / CVSSP',       '2021–2024',
         'Clarity Enhancement Challenge series (CEC1, CEC2, CEC3)'],
        ['12', 'Ephraim & Malah',       'IEEE TASLP',                  '1985',
         '"Speech enhancement using a minimum mean-square error log-spectral amplitude estimator"'],
    ]
    styled_table(doc, ['#', 'Authors', 'Venue', 'Year', 'Title / Notes'], refs)

    add_page_break(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # SLIDE 21: THANK YOU
    # ══════════════════════════════════════════════════════════════════════════
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.cell(0, 0)
    set_cell_bg(cell, DARK_BLUE)
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    p.paragraph_format.space_after  = Pt(10)
    r1 = p.add_run('Thank You\n')
    r1.bold = True; r1.font.size = Pt(36); r1.font.color.rgb = WHITE; r1.font.name = 'Calibri'
    r2 = p.add_run('We welcome your questions, feedback, and guidance.\n\n')
    r2.font.size = Pt(12); r2.font.color.rgb = DIM_TEXT; r2.font.name = 'Calibri'
    r3 = p.add_run('Jwanil Modi  23BIT194   |   Namya Shah  23BIT027\n')
    r3.bold = True; r3.font.size = Pt(12); r3.font.color.rgb = LIGHT_GRAY; r3.font.name = 'Calibri'
    r4 = p.add_run('github.com/Jwanil  ·  All code will be open source\n ')
    r4.font.size = Pt(10); r4.font.color.rgb = DIM_TEXT; r4.font.name = 'Calibri'

    doc.add_paragraph()
    add_styled_paragraph(doc,
        'Minor Project Proposal · Department of Information Technology · 2026',
        'small', color=DIM_TEXT, align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=6, space_after=4)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# ENTRY POINT
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    out_path = os.path.abspath(os.path.join(out_dir, 'Project_Proposal_Hearing_Aid_Speech_Enhancement.docx'))

    print('Building document...')
    doc = build_document()
    doc.save(out_path)
    print(f'✅ Saved: {out_path}')
    print(f'   Slides: 21  |  Sections: Cover + 20 content slides')
